#!/usr/bin/env python3
"""Build the editable template and the two source DOCX files for the kit."""

from __future__ import annotations

import argparse
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Mm, Pt, RGBColor


CONFIG = json.loads(Path(__file__).with_name("kit_config.json").read_text(encoding="utf-8"))
VERSION = CONFIG["version"]
PUBLICATION_DATE = CONFIG["publicationDate"]
PUBLICATION_DATE_LABEL = CONFIG["publicationDateLabel"]

ORGANISATION = "Hagnéré Code"
SITE = "https://hagnere-code.ai"
PAGE_DXA = 9360
LANDSCAPE_DXA = 15137
FONT = "Calibri"
INK = "1F2937"
MUTED = "5B6472"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
VIOLET = "6D28D9"
EMERALD = "047857"
AMBER = "7A5A00"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHTER = "F7F8FA"
LIGHT_GREEN = "EAF7F2"
LIGHT_AMBER = "FFF7E6"
WHITE = "FFFFFF"


@dataclass(frozen=True)
class Field:
    label: str
    prompt: str
    example: str


@dataclass(frozen=True)
class SectionSpec:
    number: int
    title: str
    why: str
    fields: Sequence[Field]
    checks: Sequence[str]
    warning: str | None = None


SECTIONS = [
    SectionSpec(
        1,
        "La décision en une page",
        "Cette synthèse évite que quinze pages de détails cachent encore la décision à prendre.",
        [
            Field("Problème métier", "Décrivez la situation à corriger en trois phrases maximum.", "Les responsables planifient les interventions dans un classeur partagé, les techniciens envoient leurs comptes rendus par messagerie et l'administration ressaisit les informations pour préparer la facturation."),
            Field("Résultat attendu", "Nommez le résultat métier observable, sans imposer de technologie.", "Disposer d'un dossier d'intervention unique, exploitable au bureau et sur mobile, depuis la planification jusqu'à la validation du rapport."),
            Field("Décision recherchée", "Que devra permettre la consultation ?", "Comparer une solution existante, une plateforme configurable et un développement spécifique sur le même périmètre et sur trois ans."),
            Field("Première version utile", "Quel résultat minimal doit fonctionner de bout en bout ?", "Planifier une intervention, l'affecter, la réaliser, joindre des photos, faire valider le rapport et exporter les éléments nécessaires à la facturation."),
            Field("Hors périmètre", "Écrivez ce qui ne doit pas être chiffré en première version.", "Comptabilité générale, paie, portail client, géolocalisation permanente et maintenance prédictive."),
        ],
        [
            "Le problème tient en trois phrases et ne présuppose pas la solution.",
            "La première version traverse un processus complet, pas une liste d'écrans.",
            "Les exclusions sont suffisamment précises pour rendre les offres comparables.",
        ],
    ),
    SectionSpec(
        2,
        "Le processus actuel",
        "Le prestataire doit comprendre le travail réel, y compris les contournements, avant de proposer des écrans.",
        [
            Field("Déclencheur", "Quel événement démarre le processus ?", "Une demande d'intervention est reçue par email ou téléphone."),
            Field("Étapes", "Décrivez chaque étape, son responsable et son support actuel.", "1. L'assistante crée une ligne dans Excel. 2. Le responsable affecte un technicien. 3. Le technicien reçoit un PDF. 4. Il renvoie photos et notes. 5. L'assistante ressaisit le rapport. 6. Le responsable valide avant facturation."),
            Field("Exceptions", "Listez les cas qui cassent le déroulé normal.", "Intervention urgente sans numéro de commande, matériel inaccessible, deuxième visite nécessaire, refus de signature, absence de réseau."),
            Field("Volumes", "Donnez des ordres de grandeur et leur saisonnalité.", "Environ 420 interventions par mois, avec un pic estimé à 650 entre novembre et février."),
            Field("Coût actuel", "Mesurez le temps, les erreurs et les délais ; indiquez les hypothèses.", "Hypothèse à vérifier pendant deux semaines : 10 heures de coordination et ressaisie par semaine, plus 6 rapports incomplets par mois."),
        ],
        [
            "Chaque étape possède un acteur, une entrée et une sortie.",
            "Les exceptions et la solution de secours sont décrites.",
            "Les volumes proviennent d'une mesure ou sont signalés comme hypothèses.",
        ],
    ),
    SectionSpec(
        3,
        "Les utilisateurs, rôles et situations d'usage",
        "Un même écran ne convient pas à une personne au bureau, à un technicien sur mobile et à un administrateur.",
        [
            Field("Profils", "Qui utilise, administre ou consulte l'outil ?", "2 planificateurs, 12 techniciens, 3 responsables d'agence, 1 administratrice et 1 prestataire de maintenance."),
            Field("Droits", "Que peut voir, créer, modifier, exporter ou supprimer chaque profil ?", "Le technicien voit uniquement ses interventions et ne supprime rien ; le planificateur affecte et corrige avant validation ; le responsable valide et exporte ; l'administratrice gère les comptes."),
            Field("Contexte", "Appareil, réseau, environnement et contraintes physiques.", "Techniciens sur smartphones récents, parfois en sous-sol sans réseau ; planification sur écrans de bureau ; gants possibles lors de la saisie."),
            Field("Accès externe", "Des clients, partenaires ou sous-traitants accèdent-ils à l'outil ?", "Aucun accès externe en première version. Les rapports validés sont envoyés en PDF."),
        ],
        [
            "Chaque profil a des droits explicites et respecte le principe du besoin d'en connaître.",
            "Les situations dégradées sont couvertes : petit écran, réseau lent ou absent, erreur de saisie.",
            "La création, la modification et la suppression de comptes ont un responsable.",
        ],
    ),
    SectionSpec(
        4,
        "Les scénarios critiques",
        "Les scénarios décrivent ce que l'utilisateur doit réussir ; ils sont plus vérifiables qu'une liste de fonctionnalités.",
        [
            Field("Scénario 1", "Profil + situation + action + résultat attendu.", "En tant que planificatrice, je crée une intervention urgente, vérifie les conflits et l'affecte à un technicien disponible afin qu'il reçoive les informations nécessaires."),
            Field("Scénario 2", "Ajoutez un scénario mobile complet.", "En tant que technicien, j'ouvre mon intervention, consulte les consignes, renseigne les mesures, ajoute trois photos et une signature puis envoie le rapport, même après une coupure de réseau."),
            Field("Scénario 3", "Ajoutez un scénario de validation ou correction.", "En tant que responsable, je refuse un rapport incomplet avec un motif ; le technicien corrige uniquement les champs concernés et le nouvel envoi reste historisé."),
            Field("Scénario 4", "Ajoutez un scénario d'administration des accès.", "En tant qu'administratrice, je désactive le compte d'un technicien parti, réattribue ses interventions ouvertes et vérifie qu'il ne peut plus se connecter, sans supprimer l'historique de ses actions."),
            Field("Scénario 5", "Ajoutez un scénario de restitution ou de réversibilité.", "En tant que responsable, j'exporte les interventions, rapports, photos et journaux d'une période donnée dans des formats documentés, puis je contrôle les relations entre les fichiers avant une migration."),
            Field("Scénario d'échec", "Que doit-il se passer si un service connecté est indisponible ?", "L'intervention reste enregistrée, l'utilisateur voit un message actionnable et l'export en échec peut être relancé sans créer de doublon."),
        ],
        [
            "Chaque scénario possède un début, une fin et un résultat observable.",
            "Au moins un cas d'erreur et un cas de reprise sont couverts.",
            "Les scénarios critiques peuvent devenir des tests de recette.",
        ],
    ),
    SectionSpec(
        5,
        "Les règles métier et les cas limites",
        "Les règles cachées dans les habitudes ou les formules Excel sont souvent le vrai coût du projet.",
        [
            Field("Règles de calcul", "Formule, unités, arrondis, exceptions et responsable de validation.", "La durée facturable est arrondie au quart d'heure supérieur après validation du responsable. Les trajets ne sont pas inclus dans ce calcul."),
            Field("États et transitions", "Quels statuts existent et qui peut les changer ?", "Brouillon > planifiée > en cours > à valider > validée > exportée. Une intervention validée ne peut être modifiée sans réouverture tracée par un responsable."),
            Field("Doublons", "Comment détecter et traiter deux enregistrements proches ?", "Alerte si même site, même équipement et même créneau ; le planificateur décide de fusionner ou de conserver."),
            Field("Traçabilité", "Quelles actions doivent rester historisées ?", "Affectation, changement de date, modification des mesures, validation, réouverture et export."),
        ],
        [
            "Chaque calcul peut être refait à partir des données conservées.",
            "Les droits de changement d'état sont explicites.",
            "Les corrections après validation restent traçables.",
        ],
    ),
    SectionSpec(
        6,
        "Les données et leur reprise",
        "La migration échoue rarement parce qu'une colonne manque ; elle échoue lorsque personne ne tranche les doublons, les valeurs vides et l'historique utile.",
        [
            Field("Sources", "Listez fichiers, bases, pièces jointes et propriétaires.", "Un classeur de 8 400 interventions, un export de 3 200 équipements, 18 Go de photos sur le NAS et des référentiels clients dans le logiciel de facturation."),
            Field("Qualité", "Doublons, champs vides, formats et valeurs incohérentes.", "Le code équipement manque sur environ 12 % des lignes ; les noms de sites ne sont pas normalisés ; certains rapports n'ont qu'une photo sans texte."),
            Field("Stratégie", "Tout migrer, dossiers actifs ou archive séparée ?", "Migrer les équipements actifs et les interventions des 24 derniers mois ; conserver l'historique antérieur en archive figée et consultable."),
            Field("Validation", "Qui accepte le résultat de la reprise et sur quel échantillon ?", "La responsable administrative valide les totaux, puis 30 dossiers tirés dans chaque année et 20 cas connus comme atypiques."),
            Field("Conservation", "Durées métier ou légales à confirmer.", "À définir avec le responsable de traitement et, si nécessaire, le conseil juridique ou comptable. Le prestataire ne déduit pas seul une durée de conservation."),
        ],
        [
            "Chaque source a un propriétaire et un format d'export identifié.",
            "Les données à corriger ou exclure sont quantifiées.",
            "La recette de migration compare comptages, totaux, échantillons et pièces jointes.",
        ],
        warning="N'insérez jamais de données clients réelles, mots de passe ou clés API dans le cahier des charges partagé.",
    ),
    SectionSpec(
        7,
        "Les intégrations et les solutions de secours",
        "Une connexion non documentée peut déplacer le prix, le délai et le niveau de risque du projet.",
        [
            Field("Systèmes", "Nom, version, propriétaire, documentation et contact.", "Logiciel de facturation : version à confirmer, documentation API non reçue. Annuaire Microsoft 365 pour les comptes, faisabilité à étudier."),
            Field("Flux", "Données, direction, fréquence, volume et déclencheur.", "Après validation, exporter chaque nuit les temps et références de commande ; importer chaque semaine la liste des clients et sites actifs."),
            Field("Erreurs", "Que voit l'utilisateur et qui est alerté ?", "L'export conserve un statut, un motif lisible et un bouton de relance. L'administratrice reçoit une alerte après trois échecs."),
            Field("Secours", "Processus manuel temporaire si l'intégration ne fonctionne pas.", "Générer un CSV au format convenu, contrôlé puis importé manuellement par l'administration."),
        ],
        [
            "Chaque flux précise son sens, son déclencheur et son propriétaire.",
            "La faisabilité non vérifiée est chiffrée comme étude ou option.",
            "Un mode dégradé empêche l'arrêt complet de l'activité.",
        ],
    ),
    SectionSpec(
        8,
        "Le périmètre et les priorités",
        "Prioriser protège le budget et permet de tester l'adoption avant d'étendre la solution.",
        [
            Field("Indispensable V1", "Fonctions sans lesquelles le processus complet ne fonctionne pas.", "Comptes et rôles, clients/sites/équipements, planning, intervention mobile, photos, validation, PDF, export de facturation, historique et sauvegardes."),
            Field("Version ultérieure", "Fonctions utiles mais non nécessaires au premier résultat.", "Portail client, stock véhicule, tableaux de bord avancés et signature électronique qualifiée."),
            Field("Hors périmètre", "Exclusions fermes pour la consultation.", "Paie, comptabilité générale, CRM commercial, remplacement du logiciel de facturation et géolocalisation permanente."),
            Field("Critère d'abandon", "Dans quel cas ne faut-il pas lancer le développement ?", "Si un logiciel existant couvre au moins les scénarios critiques sans plus de deux contournements majeurs et avec un coût total acceptable sur trois ans, privilégier son paramétrage."),
        ],
        [
            "Chaque élément apparaît dans une seule catégorie.",
            "La V1 produit un résultat de bout en bout testable.",
            "Le document autorise explicitement une conclusion sans développement sur mesure.",
        ],
    ),
    SectionSpec(
        9,
        "La sécurité, la confidentialité et la continuité",
        "Le cahier des charges doit demander des preuves et des responsabilités, pas la promesse vague d'un outil sécurisé.",
        [
            Field("Données", "Catégories, sensibilité, finalités, destinataires et cycle de vie.", "Coordonnées professionnelles, photos d'équipements et signatures de réception. Aucune donnée de santé ni donnée bancaire prévue."),
            Field("Accès", "Authentification, comptes nominatifs, MFA et revue des droits.", "Comptes nominatifs ; MFA exigée pour administrateurs et responsables ; revue trimestrielle des comptes actifs."),
            Field("Hébergement", "Pays, sous-traitants et transparence attendue.", "Hébergement dans l'Union européenne ; liste des sous-traitants et localisation effective communiquées avant signature."),
            Field("Sauvegardes", "Fréquence, rétention, isolement et test de restauration.", "Sauvegarde quotidienne, rétention 30 jours ; fréquence et objectif de restauration à confirmer par le prestataire ; test de restauration documenté avant mise en production."),
            Field("Incidents", "Détection, contact, délai d'information et journal des actions.", "Point de contact nommé et procédure documentée. Les engagements contractuels doivent être proposés sans inventer un délai universel."),
        ],
        [
            "La localisation, les sous-traitants et les transferts éventuels sont connus.",
            "Les habilitations, journaux, sauvegardes et tests de restauration ont un périmètre vérifiable.",
            "Le contrat prévoit restitution, destruction et notification des incidents.",
        ],
        warning="Ce modèle aide à poser les questions. Il ne constitue ni un audit de sécurité ni un conseil RGPD personnalisé.",
    ),
    SectionSpec(
        10,
        "Les livrables, la recette et la mise en production",
        "Un livrable n'est terminé que lorsqu'une preuve permet de l'accepter ou de le refuser.",
        [
            Field("Livrables", "Listez ce qui doit être remis, dans quel format et à qui.", "Application, code source et historique Git, schéma de données, documentation d'exploitation, guide utilisateur, plan de sauvegarde/restauration, export initial et procès-verbal de recette."),
            Field("Environnements", "Préproduction, production et données de test.", "Préproduction protégée avec données fictives ; aucune copie brute de production sans procédure validée."),
            Field("Recette", "Responsables, délai, preuves, retest et décision.", "Deux référents métier exécutent 32 scénarios. Chaque anomalie a une sévérité, une capture ou trace, un responsable et un statut de retest."),
            Field("Bascule", "Fenêtre, sauvegarde, critères go/no-go et retour arrière.", "Bascule un lundi matin après sauvegarde vérifiée. Retour à l'ancien processus si les scénarios de création, intervention ou export critique échouent."),
            Field("Stabilisation", "Durée, périmètre et canal de support.", "Période à proposer au devis avec distinction claire entre garantie corrective, assistance et évolution."),
        ],
        [
            "Les critères d'acceptation ont été rédigés avant le développement.",
            "Le plan de bascule possède une décision go/no-go et un retour arrière.",
            "Le code, la documentation et les accès figurent dans les livrables.",
        ],
    ),
    SectionSpec(
        11,
        "La gouvernance, le calendrier et le budget",
        "Le planning dépend autant des validations côté client que de la production du prestataire.",
        [
            Field("Rôles", "Sponsor, décideur, référent métier, référent données et responsable recette.", "Sponsor : direction générale. Décision produit : responsable d'exploitation. Données : responsable administrative. Recette : deux techniciens et une planificatrice."),
            Field("Disponibilité", "Temps réservé par semaine côté entreprise.", "Deux ateliers de 90 minutes au démarrage, puis 4 heures par semaine réparties entre le référent métier et les utilisateurs testeurs."),
            Field("Jalons", "Livrable, date cible, responsable de validation et délai de retour.", "Cadrage, prototype des scénarios critiques, première version, reprise à blanc, recette, formation, bascule et stabilisation. Dates à construire avec les candidats."),
            Field("Budget", "Construction, coûts récurrents, temps interne et réserve.", "Enveloppe illustrative fictive : 30 000 à 45 000 € HT pour la construction. Les licences, l'hébergement, la maintenance, les évolutions, la reprise et le temps interne sont présentés séparément."),
            Field("Comparaison", "Horizon et hypothèses communes aux offres.", "Comparer sur 36 mois avec 18 utilisateurs, 420 interventions mensuelles, le même périmètre V1 et les mêmes exigences de reprise, support et réversibilité."),
        ],
        [
            "Un seul décideur arbitre le périmètre au quotidien.",
            "Les délais de validation côté entreprise sont visibles.",
            "Chaque candidat sépare coût initial, récurrent, optionnel, interne et coût de sortie.",
        ],
    ),
    SectionSpec(
        12,
        "La maintenance, la propriété et la réversibilité",
        "La liberté de changer de prestataire dépend du contrat, mais aussi des accès, formats, documents et compétences remis.",
        [
            Field("Maintenance", "Correctif, sécurité, support et évolution : qui fait quoi ?", "Le candidat détaille heures ou forfait, plages de support, sévérités, délais proposés, exclusions, dépendances et processus de demande d'évolution."),
            Field("Propriété", "Données, code spécifique, composants tiers, comptes et licences.", "Les données restent exportables ; la cession ou licence du code spécifique est écrite au contrat ; chaque composant tiers est identifié avec sa licence."),
            Field("Accès", "Comptes qui doivent être détenus par l'entreprise.", "Organisation Git, hébergement, domaine éventuel, supervision et coffre de secrets créés au nom de l'entreprise ou avec un administrateur interne."),
            Field("Sortie", "Format, délai, coût, assistance et preuve d'importabilité.", "Export documenté des données et pièces jointes, code et procédure de redéploiement. Un test d'export puis de réimport sur un échantillon fait partie de la recette."),
            Field("Fin de vie", "Comment fermer proprement le service ?", "Planifier export final, période de lecture seule, révocation des comptes, restitution ou destruction attestée des copies et archivage de la documentation."),
        ],
        [
            "Les formats de restitution sont structurés, documentés et testés.",
            "Les accès stratégiques ne dépendent pas d'un compte personnel du prestataire.",
            "La fin de contrat prévoit restitution, assistance et destruction des copies selon le périmètre applicable.",
        ],
    ),
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        node = OxmlElement("w:tblHeader")
        node.set(qn("w:val"), "true")
        tr_pr.append(node)


def set_table_geometry(
    table,
    widths: Sequence[int],
    indent: int = 120,
    page_dxa: int = PAGE_DXA,
    cell_margin: int = 120,
) -> None:
    if sum(widths) != page_dxa:
        raise ValueError(f"Column widths must total {page_dxa}, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(page_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(
                cell,
                top=cell_margin,
                start=cell_margin,
                bottom=cell_margin,
                end=cell_margin,
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def weighted_widths(weights: Sequence[int], total: int = LANDSCAPE_DXA) -> list[int]:
    weight_total = sum(weights)
    widths = [round(total * weight / weight_total) for weight in weights]
    widths[-1] += total - sum(widths)
    return widths


def set_run_font(run, size=11, color=INK, bold=False, italic=False) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold
    run.font.italic = italic
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)


def set_style_font(style, size: float, color: str, bold=False) -> None:
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)


def configure_section_layout(section, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
        section.top_margin = Mm(15)
        section.right_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(15)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        side_margin = (Mm(210) - Inches(6.5)) // 2
        section.top_margin = Mm(22.5)
        section.right_margin = side_margin
        section.bottom_margin = Mm(22.5)
        section.left_margin = side_margin
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)


def configure_document(doc: Document, title: str, subject: str) -> None:
    configure_section_layout(doc.sections[0])

    normal = doc.styles["Normal"]
    set_style_font(normal, 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title_style = doc.styles["Title"]
    set_style_font(title_style, 26, INK, True)
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, 12, MUTED)
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, BLUE_DARK, 10, 5),
    ):
        style = doc.styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = ORGANISATION
    props.category = "Ressource opérationnelle"
    props.keywords = "cahier des charges, application métier, logiciel sur mesure, recette, réversibilité"
    props.comments = "Ressource pédagogique. Ne constitue pas un conseil juridique ou RGPD personnalisé."
    props.version = VERSION

    for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        r_pr = doc.styles[style_name].element.get_or_add_rPr()
        lang = r_pr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            r_pr.append(lang)
        lang.set(qn("w:val"), "fr-FR")

    add_header_footer(doc)


def add_field_code(paragraph, code: str, display: str) -> None:
    run = paragraph.add_run()
    for node_type, value in (("begin", None), (None, code), ("separate", None), (None, display), ("end", None)):
        if node_type:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), node_type)
        elif value == code:
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = value
        else:
            node = OxmlElement("w:t")
            node.text = value
        run._r.append(node)
    set_run_font(run, 8, MUTED)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = f"HAGNÉRÉ CODE · KIT APPLICATION MÉTIER · V{VERSION}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.runs[0], 8, MUTED, True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.add_run("hagnere-code.ai · ")
    add_field_code(footer, "PAGE", "1")
    footer.add_run(" / ")
    add_field_code(footer, "NUMPAGES", "1")
    for run in footer.runs:
        set_run_font(run, 8, MUTED)


def add_hyperlink(paragraph, label: str, url: str) -> None:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([fonts, color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_callout(doc: Document, title: str, text: str, variant="blue") -> None:
    palettes = {
        "blue": (LIGHT_BLUE, BLUE_DARK),
        "green": (LIGHT_GREEN, EMERALD),
        "amber": (LIGHT_AMBER, AMBER),
        "gray": (LIGHTER, MUTED),
    }
    fill, color = palettes[variant]
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.2
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    run = paragraph.add_run(f"{title}  ")
    set_run_font(run, 10.5, color, True)
    set_run_font(paragraph.add_run(text), 10.5, INK)


def add_cover(doc: Document, filled: bool, readme=False) -> None:
    for _ in range(4 if not readme else 2):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(kicker.add_run("KIT PRATIQUE" if not readme else "MODE D'EMPLOI"), 9, VIOLET, True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(
        "Exemple rempli - cahier des charges d'une application métier"
        if filled
        else "Cahier des charges d'une application métier"
    )
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Cas pédagogique entièrement fictif - aucune entreprise ni performance réelle"
        if filled
        else "Modèle éditable pour cadrer, comparer et préparer la recette"
    )
    metadata = doc.add_table(rows=1, cols=2)
    set_table_geometry(metadata, [2900, 6460])
    metadata.rows[0].cells[0].text = "Information"
    metadata.rows[0].cells[1].text = "Valeur"
    for cell in metadata.rows[0].cells:
        shade(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 8.5, BLUE_DARK, True)
    mark_header_row(metadata.rows[0])
    values = [
        ("Version", f"{VERSION} - {PUBLICATION_DATE_LABEL}"),
        ("Format", "Exemple de niveau de précision" if filled else "Document de travail à adapter"),
        ("Public", "Dirigeants, indépendants et responsables de projet"),
        ("Finalité", "Produire un périmètre comparable et des critères de réception"),
    ]
    for label, value in values:
        left, right = metadata.add_row().cells
        set_table_geometry(metadata, [2900, 6460])
        shade(left, LIGHT_BLUE)
        set_run_font(left.paragraphs[0].add_run(label), 9, BLUE_DARK, True)
        set_run_font(right.paragraphs[0].add_run(value), 9.5, INK)
    add_callout(
        doc,
        "Point de départ",
        "Décrivez le problème, les utilisateurs et les scénarios avant de choisir une technologie. Le modèle peut conclure qu'il faut améliorer l'existant ou acheter une solution du marché.",
        "green",
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run("Ressource gratuite - téléchargement sans email - hagnere-code.ai"), 9, MUTED)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_intro(doc: Document, filled: bool) -> None:
    doc.add_heading("Comment utiliser ce document", level=1)
    text = (
        "L'exemple montre un niveau de précision possible pour une PME fictive de maintenance. "
        "Tous les noms, volumes, montants, délais et objectifs sont inventés pour la pédagogie."
        if filled
        else "Complétez d'abord les sections 1, 2, 4 et 8. Supprimez les rubriques inutiles, conservez les inconnues avec la mention « à confirmer » et demandez aux candidats de répondre avec les mêmes hypothèses."
    )
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.keep_together = True
    add_callout(
        doc,
        "Structure du kit",
        "Le modèle décline les sept livrables du guide en quatorze rubriques guidées, puis fournit six matrices éditables pour les scénarios, les données, les droits, les intégrations, la recette et les responsabilités.",
        "blue",
    )
    steps = [
        ("1. Décrire", "Le processus réel, ses acteurs, ses exceptions et son coût actuel."),
        ("2. Arbitrer", "La première version utile, ce qui attendra et ce qui est exclu."),
        ("3. Vérifier", "Les critères de recette, les preuves, les responsabilités et le plan de sortie."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2200, 7160])
    table.rows[0].cells[0].text = "Étape"
    table.rows[0].cells[1].text = "Résultat"
    for cell in table.rows[0].cells:
        shade(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 8.5, BLUE_DARK, True)
    mark_header_row(table.rows[0])
    for label, detail in steps:
        row = table.add_row()
        set_table_geometry(table, [2200, 7160])
        shade(row.cells[0], LIGHT_BLUE)
        set_run_font(row.cells[0].paragraphs[0].add_run(label), 10, BLUE_DARK, True)
        set_run_font(row.cells[1].paragraphs[0].add_run(detail), 10, INK)
    add_callout(
        doc,
        "Confidentialité",
        "Ne collez ni données clients réelles, ni mots de passe, ni clés API dans un document transmis à plusieurs candidats. Utilisez des exemples fictifs ou anonymisés et fournissez les accès dans un canal séparé après contractualisation.",
        "amber",
    )
    doc.add_heading("Glossaire minimum", level=2)
    glossary = [
        ("API", "Interface documentée permettant à deux logiciels d'échanger des données."),
        ("MFA", "Authentification multifacteur : un second facteur en plus du mot de passe."),
        ("Recette", "Vérification formelle des critères prévus avant d'accepter un livrable."),
        ("Réversibilité", "Capacité à récupérer données, code, documents et accès pour changer de solution ou de prestataire."),
        ("TCO", "Coût total de possession : construction ou licences, temps interne, exploitation, maintenance et sortie sur une durée donnée."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2000, 7360])
    table.rows[0].cells[0].text = "Terme"
    table.rows[0].cells[1].text = "Définition"
    for cell in table.rows[0].cells:
        shade(cell, LIGHT_GRAY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 8.5, INK, True)
    mark_header_row(table.rows[0])
    for term, definition in glossary:
        row = table.add_row()
        set_table_geometry(table, [2000, 7360])
        shade(row.cells[0], LIGHT_GRAY)
        set_run_font(row.cells[0].paragraphs[0].add_run(term), 9.5, INK, True)
        set_run_font(row.cells[1].paragraphs[0].add_run(definition), 9.5, INK)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_section(doc: Document, spec: SectionSpec, filled: bool) -> None:
    doc.add_heading(f"{spec.number}. {spec.title}", level=1)
    add_callout(doc, "Pourquoi cette rubrique", spec.why, "gray")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Question à traiter"
    headers[1].text = "Votre réponse" if not filled else "Réponse du cas fictif"
    for cell in headers:
        shade(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 9, BLUE_DARK, True)
    mark_header_row(table.rows[0])
    for field in spec.fields:
        row = table.add_row()
        p = row.cells[0].paragraphs[0]
        set_run_font(p.add_run(field.label), 9.5, INK, True)
        p2 = row.cells[0].add_paragraph(field.prompt)
        set_run_font(p2.runs[0], 8.5, MUTED)
        answer = field.example if filled else f"[À compléter - {field.prompt}]\n\n"
        set_run_font(row.cells[1].paragraphs[0].add_run(answer), 9.5, INK)
        if not filled:
            shade(row.cells[1], LIGHTER)
    set_table_geometry(table, [3100, 6260])
    doc.add_heading("Terminé lorsque", level=2)
    checklist = doc.add_table(rows=1, cols=2)
    checklist.rows[0].cells[0].text = "État"
    checklist.rows[0].cells[1].text = "Critère de fin"
    for cell in checklist.rows[0].cells:
        shade(cell, LIGHT_GREEN)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 8.5, EMERALD, True)
    mark_header_row(checklist.rows[0])
    for check in spec.checks:
        row = checklist.add_row()
        set_table_geometry(checklist, [600, 8760])
        shade(row.cells[0], LIGHT_GREEN)
        marker = "☒" if filled else "☐"
        marker_p = row.cells[0].paragraphs[0]
        marker_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(marker_p.add_run(marker), 12, EMERALD, True)
        set_run_font(row.cells[1].paragraphs[0].add_run(check), 9.5, INK)
    if spec.warning:
        add_callout(doc, "Attention", spec.warning, "amber")


def add_landscape_matrix(
    doc: Document,
    title: str,
    purpose: str,
    headers: Sequence[str],
    weights: Sequence[int],
    row_keys: Sequence[str],
    example_rows: Sequence[Sequence[str]],
    filled: bool,
) -> None:
    doc.add_heading(title, level=2)
    paragraph = doc.add_paragraph(purpose)
    paragraph.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = label
        shade(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 7.2, BLUE_DARK, True)
    mark_header_row(table.rows[0])

    rows = example_rows if filled else [
        [key, *("[À compléter]" for _ in headers[1:])] for key in row_keys
    ]
    for values in rows:
        row = table.add_row()
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            if index == 0:
                shade(cell, LIGHT_GRAY)
            set_run_font(cell.paragraphs[0].add_run(value), 7.1, INK, index == 0)

    set_table_geometry(
        table,
        weighted_widths(weights),
        indent=0,
        page_dxa=LANDSCAPE_DXA,
        cell_margin=45,
    )
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_working_matrices(doc: Document, filled: bool) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section_layout(section, landscape=True)
    doc.add_heading("Annexes éditables - les six matrices de travail", level=1)
    add_callout(
        doc,
        "Mode d'emploi",
        "Ces annexes transforment les rubriques du dossier en feuilles de travail comparables. Conservez les inconnues avec la mention « à confirmer » et anonymisez les données avant tout partage.",
        "blue",
    )

    matrices = [
        (
            "A1. Scénarios métier - 5 prioritaires et 3 optionnels",
            "Les cinq premières lignes couvrent la version utile de bout en bout. Les trois suivantes servent aux variantes ou risques secondaires ; elles peuvent rester hors première version.",
            ("ID", "Acteur et but", "Déclencheur et préconditions", "Parcours nominal", "Exceptions", "Données", "Résultat", "Preuve"),
            (8, 18, 20, 24, 18, 15, 16, 16),
            tuple(f"S-{index:02d}" for index in range(1, 9)),
            (
                ("S-01", "Planificatrice - affecter une urgence", "Demande reçue ; client et site connus", "Créer, contrôler les conflits, affecter, notifier", "Aucun technicien ; priorité contradictoire", "Client, site, créneau, priorité", "Intervention affectée et visible", "Statut, notification et journal"),
                ("S-02", "Technicien - rendre le rapport", "Intervention affectée ; mobile autorisé", "Ouvrir, saisir, joindre 3 photos, signer, envoyer", "Réseau coupé ; photo trop lourde", "Mesures, photos, signature, heure", "Rapport transmis une seule fois", "Test hors ligne puis synchronisation"),
                ("S-03", "Responsable - demander une correction", "Rapport soumis mais incomplet", "Refuser avec motif, corriger, renvoyer", "Technicien absent ; deuxième refus", "Motif, champs corrigés, versions", "Nouveau rapport historisé", "Ancienne et nouvelle version"),
                ("S-04", "Responsable - valider une pièce", "Pièce à 650 € HT ; seuil fictif 500 €", "Contrôler, valider, clôturer, exporter", "Auto-validation ; justificatif absent", "Référence, montant, justificatif", "Décision traçable avant export", "Journal des deux acteurs"),
                ("S-05", "Administratrice - fermer un accès", "Départ confirmé ; compte et dossiers connus", "Désactiver, réattribuer, tester le refus", "Dossiers sans repreneur", "Compte, rôles, dossiers ouverts", "Accès coupé sans perdre l'historique", "Test de connexion et réattribution"),
                ("S-06", "Responsable - restituer les données", "Période et formats convenus", "Exporter, contrôler les liens, documenter", "Fichier manquant ; relation rompue", "Interventions, rapports, photos, journaux", "Ensemble importable et expliqué", "Contrôle des volumes et échantillon"),
                ("S-07", "Planificatrice - traiter un doublon", "Même équipement et même créneau", "Alerter, comparer, fusionner ou conserver", "Informations divergentes", "Identifiants, site, horaire", "Décision sans perte de données", "Journal de rapprochement"),
                ("S-08", "Comptabilité - relancer un export", "Service comptable indisponible", "Conserver, alerter, relancer avec même identifiant", "Deux relances simultanées", "Payload, identifiant, statut", "Un seul enregistrement créé", "Journal technique et recherche cible"),
            ),
        ),
        (
            "A2. Dictionnaire des données",
            "Une ligne par donnée importante. La durée de conservation doit être justifiée selon la finalité et les obligations applicables, pas copiée depuis l'exemple.",
            ("ID / donnée", "Source de vérité", "Format et qualité", "Finalité", "Accès", "Conservation", "Sort final"),
            (14, 18, 20, 18, 16, 18, 16),
            ("D-01", "D-02", "D-03", "D-04", "D-05", "D-06"),
            (
                ("D-01 / intervention", "Application métier", "Identifiant unique ; site obligatoire", "Planifier et tracer l'exécution", "Équipe affectée et responsables", "À définir selon obligations et activité", "Export puis suppression contrôlée"),
                ("D-02 / client", "Outil de gestion client", "Identifiant externe stable", "Rattacher site et facturation", "Planification, responsables, comptabilité", "À justifier par finalité", "Synchronisation ou archivage"),
                ("D-03 / photo", "Application métier", "JPEG ; taille et quantité limitées", "Prouver l'état et l'intervention", "Technicien affecté et responsables", "À définir selon contrat et litiges", "Export puis suppression"),
                ("D-04 / signature", "Application métier", "Image et horodatage", "Attester la remise du rapport", "Responsables et comptabilité", "À analyser juridiquement", "Restitution et purge"),
                ("D-05 / compte", "Annuaire de l'entreprise", "Email professionnel unique", "Authentifier et attribuer les droits", "Administratrice", "Durée du compte puis trace nécessaire", "Désactivation et purge ciblée"),
                ("D-06 / journal", "Application métier", "Événement, acteur, date, objet", "Sécurité, preuve et diagnostic", "Administratrice et personnes habilitées", "Durée proportionnée au risque", "Archivage ou suppression"),
            ),
        ),
        (
            "A3. Matrice des droits",
            "Une coche autorise l'action dans le périmètre indiqué. Une case vide signifie interdit, pas « à décider plus tard ».",
            ("Rôle", "Voir", "Créer", "Modifier", "Valider", "Exporter", "Supprimer", "Périmètre", "Preuve de recette"),
            (18, 8, 8, 9, 9, 9, 9, 24, 25),
            ("Rôle 1", "Rôle 2", "Rôle 3", "Rôle 4", "Rôle 5"),
            (
                ("Technicien", "Oui", "Rapport", "Ses brouillons", "Non", "Non", "Non", "Ses interventions affectées", "URL d'une autre agence refusée"),
                ("Planificatrice", "Oui", "Intervention", "Avant validation", "Non", "Non", "Non", "Son agence", "Affectation inter-agence refusée"),
                ("Responsable", "Oui", "Non", "Réouverture tracée", "Oui", "Oui", "Non", "Son agence", "Validation d'une autre agence refusée"),
                ("Comptabilité", "Validé", "Non", "Non", "Non", "Oui", "Non", "Dossiers exportables", "Brouillons invisibles"),
                ("Administratrice", "Oui", "Comptes", "Paramètres", "Non", "Technique", "Selon règle", "Toutes agences ; accès tracé", "Action sensible journalisée"),
            ),
        ),
        (
            "A4. Matrice des intégrations",
            "Chaque flux possède un propriétaire et un mode dégradé. Le mot « connecteur » ne suffit pas à prouver le comportement réel.",
            ("Flux", "Source", "Cible", "Données", "Fréquence / volume", "Authentification", "Reprise sans doublon", "Mode dégradé et responsable"),
            (14, 13, 13, 20, 16, 16, 19, 24),
            ("I-01", "I-02", "I-03", "I-04"),
            (
                ("I-01 clients", "Gestion client", "Application", "ID, nom, sites actifs", "Toutes les heures ; volume à mesurer", "Compte technique dédié", "ID source conservé ; mise à jour rejouable", "Lecture du dernier référentiel ; alerte administratrice"),
                ("I-02 export", "Application", "Comptabilité", "Intervention validée, temps, pièces", "Quotidien ; pic à mesurer", "Clé en coffre séparé", "Clé d'idempotence par intervention", "File d'attente ; relance par comptabilité"),
                ("I-03 annuaire", "Annuaire", "Application", "Compte, équipe, statut", "À l'événement", "Protocole à confirmer", "Événement versionné", "Aucun nouveau compte ; alerte administratrice"),
                ("I-04 rapports", "Application", "Messagerie", "PDF validé et destinataire", "À la validation", "Service de messagerie", "Même rapport non renvoyé sans action", "Téléchargement manuel ; responsable d'agence"),
            ),
        ),
        (
            "A5. Matrice de recette",
            "Le validateur métier décide avec une preuve conservée. Une anomalie corrigée est retestée sur le même jeu d'essai.",
            ("Test", "Données initiales", "Action", "Résultat attendu", "Preuve", "Validateur", "Sévérité si échec", "Retest"),
            (10, 19, 18, 24, 18, 15, 17, 12),
            ("R-01", "R-02", "R-03", "R-04", "R-05", "R-06"),
            (
                ("R-01", "Urgence et 2 techniciens disponibles", "Affecter puis notifier", "Un seul technicien reçoit le dossier", "Statut et notification", "Planificatrice", "Bloquante", "Date / résultat"),
                ("R-02", "Rapport avec 3 photos hors ligne", "Saisir puis reconnecter", "Tout est conservé et envoyé une fois", "Captures avant / après", "Technicien pilote", "Bloquante", "Date / résultat"),
                ("R-03", "Pièce 650 € ; seuil 500 €", "Tenter auto-validation puis faire valider", "Premier refus ; seconde décision tracée", "Journal des acteurs", "Responsable", "Majeure", "Date / résultat"),
                ("R-04", "Compte désactivé", "Ouvrir une URL connue", "Accès refusé sans effacer l'historique", "Réponse et journal", "Administratrice", "Bloquante", "Date / résultat"),
                ("R-05", "Export comptable en panne", "Relancer deux fois", "Une seule écriture cible", "ID et recherche cible", "Comptabilité", "Bloquante", "Date / résultat"),
                ("R-06", "Échantillon de sortie convenu", "Exporter puis contrôler", "Volumes et liens correspondent", "Rapport de contrôle", "Direction", "Majeure", "Date / résultat"),
            ),
        ),
        (
            "A6. Matrice des responsabilités",
            "Une activité sans décideur ni preuve reste un risque de planning. N'attribuez pas automatiquement toute décision au prestataire.",
            ("Activité", "Fournit", "Décide", "Exécute", "Valide", "Preuve attendue"),
            (22, 18, 18, 18, 18, 28),
            ("Cadrage", "Données", "Droits", "Développement", "Recette", "Bascule", "Exploitation", "Sortie"),
            (
                ("Cadrage", "Direction et utilisateurs", "Sponsor", "Chef de projet", "Décideur métier", "Périmètre et exclusions signés"),
                ("Données", "Propriétaire des données", "Métier / DPO si concerné", "Entreprise et prestataire", "Propriétaire", "Dictionnaire et rapport de qualité"),
                ("Droits", "Managers", "Direction", "Administratrice / prestataire", "Responsable sécurité", "Matrice et tests négatifs"),
                ("Développement", "Prestataire", "Arbitrages métier", "Prestataire", "Responsable projet", "Version, code et documentation"),
                ("Recette", "Prestataire et pilotes", "Métier", "Utilisateurs pilotes", "Sponsor", "Procès-verbal et anomalies"),
                ("Bascule", "Prestataire", "Direction", "Équipe projet", "Sponsor", "Plan, sauvegarde et retour arrière"),
                ("Exploitation", "Prestataire et administratrice", "Direction", "Responsable désigné", "Propriétaire du service", "Suivi incidents, coûts et accès"),
                ("Sortie", "Prestataire", "Direction", "Prestataire et entreprise", "Propriétaire des données", "Exports, accès, documentation et destruction"),
            ),
        ),
    ]

    for index, matrix in enumerate(matrices):
        add_landscape_matrix(doc, *matrix, filled=filled)
        if index < len(matrices) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section_layout(section)


def add_supplier_grid(doc: Document, filled: bool) -> None:
    doc.add_heading("13. Grille de réponse du candidat", level=1)
    paragraph = doc.add_paragraph(
        "Demandez à chaque candidat de répondre ligne par ligne. Une réponse « inclus » sans hypothèse ni preuve reste incomplète."
    )
    paragraph.paragraph_format.keep_together = True
    rows = [
        ("Compréhension du problème", "Reformulation, questions ouvertes et mauvais fits"),
        ("Périmètre V1", "Inclus, partiel, option, exclu et dépendances"),
        ("Planning", "Jalons, charge client, hypothèses et risques"),
        ("Coût initial", "Cadrage, conception, développement, reprise, recette et formation"),
        ("Coûts récurrents", "Licences, hébergement, maintenance, support et évolutions"),
        ("Données et intégrations", "Méthode, limites, étude préalable et secours"),
        ("Sécurité", "Mesures, preuves, sous-traitants et responsabilités"),
        ("Recette", "Scénarios, preuves, anomalies et retests"),
        ("Réversibilité", "Formats, code, documentation, accès, délai et coût de sortie"),
        ("Validité de l'offre", "Durée, taxes, modalités de paiement et exclusions"),
    ]
    table = doc.add_table(rows=1, cols=3)
    headers = table.rows[0].cells
    for cell, label in zip(headers, ("Sujet", "Réponse attendue", "Réponse du candidat")):
        cell.text = label
        shade(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 8.5, BLUE_DARK, True)
    mark_header_row(table.rows[0])
    for label, expected in rows:
        row = table.add_row()
        set_run_font(row.cells[0].paragraphs[0].add_run(label), 8.8, INK, True)
        set_run_font(row.cells[1].paragraphs[0].add_run(expected), 8.8, INK)
        answer = "Voir proposition annexée et matrice de couverture." if filled else "[À compléter par le candidat]\n\n"
        set_run_font(row.cells[2].paragraphs[0].add_run(answer), 8.8, INK)
        if not filled:
            shade(row.cells[2], LIGHTER)
    set_table_geometry(table, [2200, 3560, 3600])


def add_final_review(doc: Document, filled: bool) -> None:
    doc.add_heading("14. Contrôle avant envoi", level=1)
    checks = [
        "Le problème et le résultat sont décrits sans imposer une technologie.",
        "La V1, les options et les exclusions ne se chevauchent pas.",
        "Les données, intégrations et règles encore inconnues sont signalées.",
        "Les scénarios critiques deviennent des critères de recette.",
        "Les coûts sont comparés sur le même horizon et le même nombre d'utilisateurs.",
        "Les responsabilités côté entreprise et prestataire sont nommées.",
        "Les exigences de sécurité demandent une preuve et un responsable.",
        "La restitution des données, du code, des accès et de la documentation est prévue.",
        "Aucun secret ni donnée personnelle inutile ne figure dans le document partagé.",
        "La décision « ne pas développer » reste possible.",
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "État"
    table.rows[0].cells[1].text = "Point de contrôle"
    for cell in table.rows[0].cells:
        shade(cell, LIGHT_GREEN if filled else LIGHT_GRAY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 8.5, EMERALD if filled else INK, True)
    mark_header_row(table.rows[0])
    for check in checks:
        row = table.add_row()
        set_table_geometry(table, [600, 8760])
        shade(row.cells[0], LIGHT_GREEN if filled else LIGHTER)
        marker = "☒" if filled else "☐"
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(marker), 12, EMERALD if filled else MUTED, True)
        set_run_font(row.cells[1].paragraphs[0].add_run(check), 9.5, INK)
    add_callout(
        doc,
        "Ce que le kit ne remplace pas",
        "Le cahier des charges ne remplace pas le contrat, son ordre de priorité, un audit de sécurité, une analyse d'impact, ni un conseil juridique, social, fiscal ou sectoriel adapté.",
        "amber",
    )


def build_main_document(output: Path, filled: bool) -> None:
    doc = Document()
    title = (
        "Exemple rempli - cahier des charges d'une application métier"
        if filled
        else "Modèle de cahier des charges d'une application métier"
    )
    configure_document(doc, title, "Cadrage, comparaison, recette et réversibilité d'une application métier")
    add_cover(doc, filled)
    add_intro(doc, filled)
    if filled:
        add_callout(
            doc,
            "Entreprise fictive",
            "Alp'Interventions, PME de maintenance de 18 salariés située en Savoie, est un exemple inventé. Les personnes, volumes, coûts et cibles n'ont aucun lien avec un client réel.",
            "amber",
        )
    for spec in SECTIONS:
        add_section(doc, spec, filled)
    add_working_matrices(doc, filled)
    add_supplier_grid(doc, filled)
    add_final_review(doc, filled)
    doc.save(output)


def add_source(doc: Document, label: str, url: str, note: str) -> None:
    paragraph = doc.add_paragraph()
    set_run_font(paragraph.add_run(f"{label}. "), 9.5, INK, True)
    add_hyperlink(paragraph, "Consulter la source", url)
    set_run_font(paragraph.add_run(f" - {note}"), 9.5, MUTED)


def build_readme(output: Path) -> None:
    doc = Document()
    configure_document(doc, "Mode d'emploi - kit application métier", "Prise en main, limites et sources du kit")
    add_cover(doc, False, readme=True)
    doc.add_heading("1. Commencez par quatre rubriques", level=1)
    doc.add_paragraph(
        "Complétez la décision en une page, le processus actuel, les scénarios critiques et le périmètre. Vous obtenez déjà une base suffisante pour vérifier si un logiciel existant mérite d'être testé avant tout développement."
    )
    add_callout(doc, "Bon résultat", "Trois candidats lisent le même besoin, posent des questions comparables et indiquent clairement leurs hypothèses, exclusions et coûts sur 36 mois.", "green")
    doc.add_heading("2. Écrivez le besoin, pas l'écran", level=1)
    doc.add_paragraph(
        "« Un bouton rouge pour clôturer » impose une solution. « Le responsable clôture après avoir vérifié les champs obligatoires, avec un motif si le dossier est refusé » décrit un comportement testable. Une technologie imposée par l'existant peut rester dans le document, mais sa raison et ses limites doivent être explicites."
    )
    doc.add_heading("3. Transformez les scénarios en recette", level=1)
    doc.add_paragraph(
        "Chaque scénario critique devient un test avec données de départ, action, résultat attendu et preuve. Écrivez aussi les cas d'erreur, le mode dégradé, le retest et la personne autorisée à accepter le livrable."
    )
    doc.add_heading("4. Comparez à périmètre constant", level=1)
    comparison = doc.add_table(rows=1, cols=3)
    for cell, label in zip(comparison.rows[0].cells, ("À comparer", "Même hypothèse", "À isoler")):
        cell.text = label
        shade(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 9, BLUE_DARK, True)
    mark_header_row(comparison.rows[0])
    for values in (
        ("Utilisateurs", "Même nombre et croissance", "Comptes externes ou lecture seule"),
        ("Périmètre", "Même V1 et mêmes intégrations", "Options et études préalables"),
        ("Horizon", "36 mois", "Coût initial, récurrent, interne et sortie"),
        ("Support", "Même plage et sévérités", "Garantie, TMA et évolution"),
        ("Réversibilité", "Même livrables et formats", "Assistance et coût de sortie"),
    ):
        row = comparison.add_row()
        for index, value in enumerate(values):
            set_run_font(row.cells[index].paragraphs[0].add_run(value), 9, INK, index == 0)
    set_table_geometry(comparison, [2100, 3500, 3760])
    doc.add_heading("5. Partagez sans exposer vos données", level=1)
    doc.add_paragraph(
        "Remplacez les noms, coordonnées et exemples réels par des données fictives. Ne placez jamais un mot de passe, un jeton, une clé API ou un export client brut dans le cahier des charges. Les accès de test sont transmis plus tard dans un canal séparé, avec durée et destinataires définis."
    )
    add_callout(doc, "Limite", "Le modèle aide à cadrer les questions de données, sécurité et sous-traitance. Il ne prouve aucune conformité et ne remplace pas les analyses adaptées à votre activité.", "amber")
    doc.add_heading("6. Sources de méthode", level=1)
    doc.add_paragraph(
        "Ces sources officielles ont guidé les rubriques. Elles sont consultées le 20 juillet 2026. La page France Num porte sur les sites internet ; ses principes de cadrage sont transposés ici en le signalant."
    )
    add_source(doc, "France Num", "https://www.francenum.gouv.fr/guides-et-conseils/developpement-commercial/site-web/batir-le-cahier-des-charges-du-site-internet", "objectifs, comparaison des prestataires, responsabilités, budget et délais")
    add_source(doc, "DINUM - DesignGouv", "https://design.numerique.gouv.fr/bien-concevoir/", "commencer par les besoins, tester avec de vrais usagers, prioriser et itérer")
    add_source(doc, "RGESN", "https://ecoresponsable.numerique.gouv.fr/publications/referentiel-general-ecoconception/critere/1.2/", "cibles, besoins métier, attentes réelles et vérification d'une solution existante")
    add_source(doc, "CNIL - architecture", "https://www.cnil.fr/fr/faire-un-choix-eclaire-de-son-architecture", "flux et cycle de vie des données, hébergement, habilitations et portabilité")
    add_source(doc, "CNIL - sous-traitance", "https://www.cnil.fr/fr/securite-gerer-la-sous-traitance", "contrat, authentification, incidents, restitution et destruction")
    add_source(doc, "ANSSI - MonServiceSécurisé", "https://aide.monservicesecurise.cyber.gouv.fr/fr/article/mon-fournisseur-de-service-me-dit-que-cest-securise-8ldkcu/", "demander des mesures et engagements vérifiables au fournisseur")
    doc.add_heading("7. Conditions d'utilisation", level=1)
    doc.add_paragraph(
        "Vous pouvez modifier le kit pour vos propres projets et le partager avec vos équipes ou les prestataires consultés. La revente ou republication du kit comme ressource autonome n'est pas autorisée. Hagnéré Code ne garantit ni l'exhaustivité du modèle ni le résultat d'un projet mené à partir de celui-ci."
    )
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    args.output.mkdir(parents=True, exist_ok=True)
    build_main_document(args.output / "modele-cahier-des-charges-application-metier.docx", filled=False)
    build_main_document(args.output / "exemple-rempli-cahier-des-charges-application-metier-source.docx", filled=True)
    build_readme(args.output / "mode-emploi-cahier-des-charges-application-metier-source.docx")


if __name__ == "__main__":
    main()
