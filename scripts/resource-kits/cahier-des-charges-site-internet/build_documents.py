#!/usr/bin/env python3
"""Build the public DOCX/PDF assets for the website-specification kit.

The source is deliberately versioned: public binaries can be regenerated,
reviewed and corrected without editing opaque Office files by hand.
"""

from __future__ import annotations

import argparse
import json
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Mm, Pt, RGBColor


CONFIG = json.loads(
    Path(__file__).with_name("kit_config.json").read_text(encoding="utf-8")
)
VERSION = CONFIG["version"]
PUBLICATION_DATE = date.fromisoformat(CONFIG["publicationDate"])
PUBLICATION_DATE_LABEL = CONFIG["publicationDateLabel"]
ORGANISATION = "Hagnéré Code"
ACCENT = "2E74B5"
ACCENT_DARK = "1E4F7A"
VIOLET = "6D28D9"
EMERALD = "047857"
TEXT = "1F2937"
MUTED = "5B6472"
LIGHT = "E8EEF5"
LIGHTER = "F6F8FB"
AMBER_BG = "FFF7E6"
AMBER = "9A6700"
WHITE = "FFFFFF"
PAGE_DXA = 9360


@dataclass(frozen=True)
class Field:
    label: str
    prompt: str
    example: str


@dataclass(frozen=True)
class SectionSpec:
    number: int
    title: str
    tag: str
    why: str
    fields: Sequence[Field]
    provider: Sequence[str]
    done: Sequence[str]
    note: str | None = None


def pick(filled: bool, field: Field) -> str:
    return field.example if filled else f"[À compléter — {field.prompt}]"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int], indent: int = 120) -> None:
    if sum(widths) != PAGE_DXA:
        raise ValueError(f"Column widths must total {PAGE_DXA}, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        node = OxmlElement("w:tblHeader")
        node.set(qn("w:val"), "true")
        tr_pr.append(node)


def set_repeat_table_header(table) -> None:
    mark_header_row(table.rows[0])


def set_paragraph_keep(paragraph, keep_with_next=False, keep_together=False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_with_next and p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_together and p_pr.find(qn("w:keepLines")) is None:
        p_pr.append(OxmlElement("w:keepLines"))


def add_field_code(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def set_doc_language(doc: Document, lang: str = "fr-FR") -> None:
    # Word stores language in several independent places. Setting only the
    # visible styles leaves themeFontLang/docDefaults in en-US, which causes
    # LibreOffice to export a French PDF whose catalogue still says /Lang
    # (en-US). Set the package-wide defaults as well as the visible styles.
    styles_root = doc.styles.element
    doc_defaults = styles_root.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_root.insert(0, doc_defaults)
    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(r_pr_default)
    default_r_pr = r_pr_default.find(qn("w:rPr"))
    if default_r_pr is None:
        default_r_pr = OxmlElement("w:rPr")
        r_pr_default.append(default_r_pr)
    default_lang = default_r_pr.find(qn("w:lang"))
    if default_lang is None:
        default_lang = OxmlElement("w:lang")
        default_r_pr.append(default_lang)
    for attribute in ("val", "eastAsia", "bidi"):
        default_lang.set(qn(f"w:{attribute}"), lang)

    settings_root = doc.settings.element
    theme_lang = settings_root.find(qn("w:themeFontLang"))
    if theme_lang is None:
        theme_lang = OxmlElement("w:themeFontLang")
        settings_root.append(theme_lang)
    for attribute in ("val", "eastAsia", "bidi"):
        theme_lang.set(qn(f"w:{attribute}"), lang)

    styles = doc.styles
    for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        r_pr = style.element.get_or_add_rPr()
        lang_el = r_pr.find(qn("w:lang"))
        if lang_el is None:
            lang_el = OxmlElement("w:lang")
            r_pr.append(lang_el)
        lang_el.set(qn("w:val"), lang)
        lang_el.set(qn("w:eastAsia"), lang)
        lang_el.set(qn("w:bidi"), lang)


def set_style_font(style, name: str, size: float, color: str, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11, TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    set_style_font(title, "Calibri", 23, TEXT, True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, "Calibri", 11, MUTED)
    subtitle.paragraph_format.space_after = Pt(12)

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, "Calibri", 16, ACCENT, True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, "Calibri", 13, ACCENT_DARK, True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, "Calibri", 12, ACCENT_DARK, True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        set_style_font(style, "Calibri", 11, TEXT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.187)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_page(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = True
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    # 22.45 mm leaves the same 165.1 mm / 9 360 dxa content width as the
    # original layout, so every audited table keeps its exact geometry.
    section.left_margin = Mm(22.45)
    section.right_margin = Mm(22.45)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)


def configure_properties(doc: Document, title: str, subject: str) -> None:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = ORGANISATION
    props.keywords = "cahier des charges, site internet, modèle, recette, TPE, PME"
    props.comments = (
        "Ressource pédagogique. Ne constitue pas un conseil juridique personnalisé."
    )
    props.category = "Ressource opérationnelle"
    props.version = VERSION


def add_header_footer(doc: Document, filled: bool, label_override: str | None = None) -> None:
    for section in doc.sections:
        label = label_override or ("CAS FICTIF · " if filled else "MODÈLE ÉDITABLE · ")
        for header in (section.header, section.even_page_header):
            p = header.paragraphs[0]
            p.text = f"{ORGANISATION.upper()} · {label}v{VERSION} · {PUBLICATION_DATE.strftime('%d/%m/%Y')}"
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.08)
            p.paragraph_format.right_indent = Inches(0.08)
            p_pr = p._p.get_or_add_pPr()
            p_shd = OxmlElement("w:shd")
            p_shd.set(qn("w:fill"), TEXT)
            p_pr.append(p_shd)
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor.from_string(WHITE)
            p.runs[0].font.size = Pt(8)

        for footer in (section.footer, section.even_page_footer):
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("Hagnéré Code · Kit cahier des charges · Page ")
            run.font.color.rgb = RGBColor.from_string(MUTED)
            run.font.size = Pt(8)
            add_field_code(p, "PAGE")
            p.add_run(" / ")
            add_field_code(p, "NUMPAGES")
            p.add_run(" · hagnere-code.ai")
            for item in p.runs:
                item.font.color.rgb = RGBColor.from_string(MUTED)
                item.font.size = Pt(8)


def add_tag(paragraph, label: str, color: str = VIOLET) -> None:
    run = paragraph.add_run(label.upper())
    run.font.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(color)
    set_paragraph_keep(paragraph, keep_with_next=True, keep_together=True)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = paragraph.add_run(bold_prefix)
        r1.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    set_paragraph_keep(paragraph, keep_together=True)


def add_bullets(
    doc: Document,
    items: Iterable[str],
    numbered: bool = False,
    compact: bool = False,
) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(item, style=style)
        if compact:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            for run in p.runs:
                run.font.size = Pt(8.5)
        set_paragraph_keep(p, keep_together=True)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    """Append a real external hyperlink instead of printing a fragile URL."""
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([properties, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_link_bullets(doc: Document, items: Iterable[tuple[str, str]]) -> None:
    for label, url in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        add_hyperlink(paragraph, label, url)
        set_paragraph_keep(paragraph, keep_together=True)


def add_callout(doc: Document, title: str, text: str, variant: str = "blue") -> None:
    palette = {
        "blue": (LIGHT, ACCENT_DARK),
        "amber": (AMBER_BG, AMBER),
        "emerald": ("EAF7F2", EMERALD),
        "gray": (LIGHTER, MUTED),
    }
    bg, color = palette[variant]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.18
    p_pr = p._p.get_or_add_pPr()
    p_shd = OxmlElement("w:shd")
    p_shd.set(qn("w:fill"), bg)
    p_pr.append(p_shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), color)
        borders.append(border)
    p_pr.append(borders)
    title_run = p.add_run(f"{title} — ")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(color)
    body_run = p.add_run(text)
    body_run.font.color.rgb = RGBColor.from_string(TEXT)
    set_paragraph_keep(p, keep_together=True)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = TEXT, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)


def add_field_table(doc: Document, fields: Sequence[Field], filled: bool) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "QUESTION / DÉCISION", bold=True, color=WHITE)
    set_cell_text(table.rows[0].cells[1], "VOTRE RÉPONSE", bold=True, color=WHITE)
    for cell in table.rows[0].cells:
        shade(cell, ACCENT_DARK)
    for field in fields:
        cells = table.add_row().cells
        set_cell_text(cells[0], field.label, bold=True, color=ACCENT_DARK)
        set_cell_text(cells[1], pick(filled, field), color=TEXT if filled else MUTED)
        shade(cells[0], LIGHTER)
    set_table_geometry(table, [2850, 6510])
    set_repeat_table_header(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_matrix(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int],
    *,
    empty_rows: int = 0,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=WHITE, size=8)
        shade(table.rows[0].cells[i], ACCENT_DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, color=TEXT, size=8)
            if len(table.rows) % 2 == 0:
                shade(cells[i], LIGHTER)
    for _ in range(empty_rows):
        cells = table.add_row().cells
        for i in range(len(headers)):
            set_cell_text(cells[i], "[À compléter]" if i == 0 else "", color=MUTED, size=8)
    set_table_geometry(table, widths)
    set_repeat_table_header(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_section_intro(doc: Document, spec: SectionSpec) -> None:
    heading = doc.add_heading(f"{spec.number}. {spec.title}", level=1)
    add_tag(heading.insert_paragraph_before(), spec.tag, EMERALD if "ESSENTIEL" in spec.tag else VIOLET)
    add_callout(doc, "Pourquoi cette section compte", spec.why, "blue")
    if spec.note:
        add_callout(doc, "Point de vigilance", spec.note, "amber")


def add_provider_and_done(doc: Document, spec: SectionSpec) -> None:
    doc.add_heading("Ce que le prestataire doit préciser", level=2)
    add_bullets(doc, spec.provider)
    doc.add_heading("La section est terminée lorsque…", level=2)
    add_bullets(doc, spec.done)


def cover(doc: Document, filled: bool) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    add_tag(p, "CAS ENTIÈREMENT FICTIF" if filled else "MODÈLE WORD ÉDITABLE", VIOLET)
    title = doc.add_paragraph(style="Title")
    title.add_run(
        "Exemple rempli — cahier des charges de site internet"
        if filled
        else "Cahier des charges de site internet"
    )
    sub = doc.add_paragraph(style="Subtitle")
    sub.add_run(
        "Entreprise, chiffres, budget, calendrier et résultats inventés uniquement pour expliquer le modèle."
        if filled
        else "Un cahier de travail pour cadrer le besoin, comparer les offres et préparer une recette vérifiable."
    )
    add_callout(
        doc,
        "À lire avant de partager",
        (
            "Ce document illustre un projet fictif. Il ne décrit aucun client de Hagnéré Code et ne constitue ni un devis ni une preuve de résultats."
            if filled
            else "Ne saisissez aucun mot de passe, clé API, donnée client brute, donnée sensible ou secret commercial détaillé. Partagez une copie expurgée et transmettez les accès séparément par un canal sécurisé."
        ),
        "amber",
    )
    identity_heading = doc.add_heading(
        "1. Identité, statut et contrôle du document",
        level=1,
    )
    add_tag(
        identity_heading.insert_paragraph_before(),
        "ESSENTIEL · CONTRÔLE",
        EMERALD,
    )
    fields = [
        Field("Entreprise", "raison sociale ou nom d'usage", "Entreprise Exemple Services — cas fictif"),
        Field("Projet", "nom court et compréhensible", "Refonte du site de génération de demandes B2B"),
        Field("Type", "création ou refonte", "Refonte d'un site WordPress créé en 2019"),
        Field("Responsable", "fonction, sans donnée personnelle inutile", "Responsable communication"),
        Field("Version / statut", "ex. v0.3 — validé en interne", f"v{VERSION} — exemple pédagogique prêt à consulter"),
        Field("Date limite de réponse", "date et fuseau si utile", "30 août 2026 — date fictive"),
        Field("Diffusion", "interne, candidats, confidentiel", "Candidats présélectionnés — copie expurgée"),
    ]
    add_field_table(doc, fields, filled)
    add_callout(
        doc,
        "Portée du document",
        "Ce cahier des charges exprime un besoin. Il peut devenir une pièce contractuelle s'il est intégré au contrat ; sa portée dépend de l'ensemble contractuel et de l'ordre de priorité des documents. Faites valider les clauses adaptées à votre situation.",
        "gray",
    )
    doc.add_page_break()


def add_usage(doc: Document, filled: bool) -> None:
    doc.add_heading("Mode d'emploi — 60 à 90 minutes pour une première version", level=1)
    add_bullets(
        doc,
        [
            "Commencez par les rubriques ESSENTIEL. Une réponse honnête « à confirmer » vaut mieux qu'une invention.",
            "Utilisez ce document comme socle de cadrage. Pour une vente en ligne, des comptes utilisateurs, plusieurs langues ou une activité réglementée, ajoutez un module dédié et faites valider les obligations propres au service avant la consultation.",
            "Supprimez les consignes et lignes inutiles avant envoi ; gardez la version source et diffusez la même version PDF à tous les candidats.",
            "Centralisez les questions et envoyez les mêmes réponses à tous sous forme d'addendum daté.",
            "Demandez aux prestataires de distinguer ce qui est inclus, partiel, optionnel, proposé autrement ou exclu.",
            "Utilisez la grille de recette du kit pour préparer les contrôles avant même de choisir le prestataire.",
        ],
        numbered=True,
    )
    add_callout(
        doc,
        "Légende",
        "ESSENTIEL = nécessaire pour comparer les offres · OPTIONNEL = à conserver si pertinent · PRESTATAIRE RÉPOND = réponse obligatoire dans l'offre · À FAIRE VALIDER = arbitrage juridique, sécurité ou métier à confirmer.",
        "emerald",
    )
    doc.add_heading("Mini-diagnostic : quels modules garder ?", level=2)
    diagnostic_rows = [
        ("Refonte d'un site existant ?", "Garder le module migration SEO et l'inventaire des URL.", "Oui"),
        ("Vente ou souscription en ligne ?", "Ajouter un module catalogue, prix, commande, paiement, livraison, retours, support et obligations B2C/B2B à valider.", "Non"),
        ("Comptes utilisateurs ou espace client ?", "Ajouter authentification, rôles, cycle de vie, récupération, journalisation, sécurité, données et tests d'accès.", "Non"),
        ("Connexion à un CRM, ERP, agenda ou paiement ?", "Garder intégrations, données de test et solution de secours.", "Oui — CRM à confirmer"),
        ("Données sensibles ou activité réglementée ?", "Escalade vers les spécialistes compétents avant chiffrage.", "Non identifié"),
        ("Plusieurs langues ?", "Ajouter langues et variantes, traduction, validation, gestion éditoriale, URL et SEO par langue.", "Non en V1"),
    ]
    add_matrix(
        doc,
        ["QUESTION", "CONSÉQUENCE", "VOTRE RÉPONSE"],
        [r if filled else (r[0], r[1], "[Oui / Non / À confirmer]") for r in diagnostic_rows],
        [2700, 4260, 2400],
    )
    doc.add_heading("Journal des versions", level=2)
    version_rows = [
        (
            VERSION,
            PUBLICATION_DATE.strftime("%d/%m/%Y"),
            "Responsable communication",
            "Version fictive prête à consultation",
        ),
    ]
    add_matrix(
        doc,
        ["VERSION", "DATE", "RESPONSABLE", "MODIFICATION"],
        version_rows if filled else [],
        [1100, 1600, 2360, 4300],
        empty_rows=2 if not filled else 0,
    )


SECTIONS = [
    SectionSpec(
        2,
        "Synthèse décisionnelle",
        "ESSENTIEL · LE DIRIGEANT DÉCIDE",
        "Une page claire évite que les candidats découvrent le vrai projet au milieu d'une annexe.",
        [
            Field("Problème métier", "ce qui ne fonctionne plus aujourd'hui", "Le site ne soutient plus la prospection : l'équipe dépend du prestataire pour chaque modification et les demandes sont peu qualifiées."),
            Field("Pourquoi maintenant", "déclencheur réel", "Préparer un salon professionnel fictif du 24 novembre 2026 et fiabiliser l'acquisition locale avant cette date."),
            Field("Résultat principal", "résultat métier, pas une fonctionnalité", "Augmenter les demandes qualifiées suivies dans le CRM ; cible illustrative : 7 à 14 par mois dans les neuf mois."),
            Field("Résultats secondaires", "deux maximum", "Préserver les pages utiles lors de la refonte ; rendre l'équipe autonome pour publier et modifier les pages courantes."),
            Field("Public prioritaire", "qui doit réussir son parcours", "Dirigeants de PME industrielles en Auvergne-Rhône-Alpes recherchant un accompagnement opérationnel."),
            Field("Action principale", "ce que le visiteur doit pouvoir faire", "Comprendre l'offre adaptée puis demander un échange en donnant le contexte minimal."),
            Field("Indispensable au lancement", "éléments non négociables", "Pages d'offres, références, ressources, formulaire qualifiant, migration, mesures essentielles et remise des accès."),
            Field("Hors périmètre", "ce que le projet ne comprend pas", "Espace client, paiement, chatbot, multilingue, application mobile et refonte complète de l'identité."),
            Field("Budget", "fourchette HT/TTC et coûts récurrents", "12 000 à 16 000 € HT pour la construction ; coûts récurrents à détailler séparément."),
            Field("Date et raison", "souhaitée ou impérative", "Mise en ligne souhaitée le 15 novembre 2026, neuf jours avant le salon ; arbitrage possible si le risque qualité augmente."),
            Field("Décideur final", "fonction et mode d'arbitrage", "Direction générale, après avis de la responsable communication."),
            Field("Projet réussi", "définition simple et observable", "Le périmètre V1 est livré et accepté, les URL utiles sont redirigées, les accès sont remis, l'équipe sait publier et la mesure des demandes fonctionne."),
        ],
        [
            "Reformuler le problème et signaler les hypothèses qui influencent prix ou délai.",
            "Indiquer explicitement toute exclusion ou alternative proposée.",
            "Séparer le prix initial, les options, les coûts récurrents et le coût de sortie.",
        ],
        [
            "Un lecteur externe comprend le projet en moins de trois minutes.",
            "Les indispensables et les exclusions ne se contredisent pas.",
            "Le budget, la date et le décideur sont explicites ou assumés « à confirmer ».",
        ],
        "Un objectif commercial ou SEO sert au pilotage ; il ne devient pas automatiquement une obligation de résultat du prestataire.",
    ),
    SectionSpec(
        3,
        "Entreprise, activité et situation actuelle",
        "ESSENTIEL",
        "Le même site n'a pas le même rôle selon le cycle de vente, la saisonnalité, l'équipe et les outils déjà en place.",
        [
            Field("Activité et offres", "décrire simplement ce qui est vendu", "Conseil B2B fictif : diagnostic opérationnel et accompagnement de transformation pour PME."),
            Field("Taille et organisation", "effectif et personnes impliquées", "12 salariés ; direction, responsable communication et deux experts contribueront aux contenus."),
            Field("Clients et zone", "segments et territoire", "PME de 20 à 250 salariés, principalement en Auvergne-Rhône-Alpes."),
            Field("Cycle de vente", "durée et étapes", "Cycle de un à trois mois : recherche, vérification des références, échange exploratoire puis proposition."),
            Field("Saisonnalité", "périodes fortes ou contraintes", "Prospection accrue avant les salons de novembre et les budgets de début d'année."),
            Field("Site actuel", "URL, technologie connue, année", "Domaine fictif : https://example.com ; WordPress créé en 2019 ; 42 URL recensées à confirmer."),
            Field("Ce qui fonctionne", "actifs à préserver", "Six pages de services reçoivent encore des visites et plusieurs références rassurent les prospects."),
            Field("Irritants observés", "faits ou retours, pas des adjectifs", "Publication lente, formulaire peu qualifiant, navigation mobile chargée, dépendance au prestataire."),
            Field("Actifs disponibles", "domaine, contenus, photos, charte, comptes", "Domaine au nom de l'entreprise, charte légère, photos avec droits à vérifier, Search Console et outil de mesure accessibles."),
            Field("Pourquoi ne pas conserver l'existant", "limite déterminante", "Le thème et les extensions rendent les évolutions risquées ; une étude de reprise reste demandée avant de conclure à une refonte complète."),
        ],
        [
            "Confirmer sa compréhension du contexte et lister les informations manquantes.",
            "Identifier les actifs réutilisables, les dépendances et les licences à vérifier.",
            "Proposer, si pertinent, une option plus simple qu'une refonte complète.",
        ],
        [
            "Les contraintes qui changent le périmètre sont visibles.",
            "Les accès existants sont listés sans inclure de secrets.",
            "Les éléments à préserver et ceux à remettre en question sont distingués.",
        ],
    ),
    SectionSpec(
        4,
        "Objectifs, situation de départ et mesure",
        "ESSENTIEL · MESURE",
        "Une cible sans point de départ ni source ne permet ni pilotage ni apprentissage.",
        [],
        [
            "Distinguer les résultats métier qu'il contribue à favoriser des critères de livraison qu'il contrôle.",
            "Décrire les événements mesurés, leur déclenchement et la procédure de vérification.",
            "Ne garantir ni chiffre d'affaires, ni volume de contacts, ni position dans un moteur de recherche.",
        ],
        [
            "Chaque objectif a une source de départ, une cible, un horizon et un responsable.",
            "Les objectifs métier et les critères de recette figurent dans deux listes différentes.",
            "La collecte de mesure est proportionnée et ses conditions juridiques restent à valider.",
        ],
        "Une cible est une hypothèse de pilotage. Le kit n'en fait pas une promesse de résultat.",
    ),
    SectionSpec(
        5,
        "Publics et parcours principaux",
        "ESSENTIEL · PÉDAGOGIE",
        "Un parcours réel aide davantage à concevoir qu'un persona décoratif rempli d'informations sans effet sur le site.",
        [
            Field("Public prioritaire", "situation, besoin, obstacle", "Dirigeant de PME qui cherche un partenaire crédible, compare plusieurs approches et manque de temps pour décoder le jargon."),
            Field("Ce qu'il doit comprendre", "message ou preuve", "Les problèmes traités, la méthode, les limites et un exemple de mission comparable."),
            Field("Ce qui le rassure", "preuves attendues", "Expertise explicitée, références vérifiables, déroulé d'une mission et contact humain identifiable."),
            Field("Contexte d'usage", "appareil, environnement, urgence", "Recherche souvent mobile après recommandation, puis comparaison sur ordinateur."),
            Field("Action attendue", "verbe et résultat", "Choisir une offre pertinente et envoyer une demande avec contexte, échéance et taille d'entreprise."),
            Field("Obstacle fréquent", "raison d'abandon", "Offres trop abstraites, formulaire trop long, absence de repères sur la suite."),
            Field("Parcours secondaire", "autre public utile", "Responsable opérations qui télécharge une ressource, la partage en interne puis revient demander un échange."),
            Field("Erreur / abandon", "ce qui doit arriver", "Conserver les informations non sensibles si erreur, expliquer le problème et proposer un contact alternatif."),
        ],
        [
            "Cartographier les étapes et états nécessaires, y compris chargement, erreur, confirmation et absence de résultat.",
            "Signaler les informations manquantes qui empêchent de concevoir un parcours.",
            "Indiquer les responsabilités de traitement après chaque demande reçue.",
        ],
        [
            "Chaque public modifie réellement un contenu, une preuve ou un parcours.",
            "Les erreurs et confirmations sont décrites, pas seulement le scénario idéal.",
            "Le responsable interne de chaque demande est identifié.",
        ],
    ),
    SectionSpec(
        6,
        "Périmètre, pages et contenus",
        "ESSENTIEL · COMPARAISON DES OFFRES",
        "Le nombre de pages ne suffit pas : chaque page doit avoir un but, un propriétaire de contenu et un statut de migration.",
        [],
        [
            "Chiffrer séparément création, reprise, réécriture, migration et saisie des contenus.",
            "Confirmer le nombre de modèles de page, les limites du CMS et la responsabilité de validation.",
            "Identifier les contenus ou droits manquants avant le démarrage.",
        ],
        [
            "Toutes les pages indispensables ont un objectif et un responsable.",
            "La V1, les options, la version ultérieure et le hors périmètre sont séparés.",
            "Les droits des textes, images, vidéos et documents sont à vérifier.",
        ],
    ),
    SectionSpec(
        7,
        "Fonctionnalités et cas limites",
        "ESSENTIEL · CHIFFRAGE",
        "Une fonctionnalité devient chiffrable lorsqu'on décrit le besoin, le comportement, les limites et la preuve attendue.",
        [],
        [
            "Répondre pour chaque ligne : inclus, partiel, optionnel, alternative proposée ou exclu.",
            "Préciser dépendances, limites de volume, coûts tiers et impact sur le calendrier.",
            "Décrire les erreurs, indisponibilités et solutions de secours prévues.",
        ],
        [
            "Chaque indispensable a un identifiant et un critère d'acceptation.",
            "Les cas d'erreur et volumes susceptibles de changer l'architecture sont décrits.",
            "Les options ne sont pas mélangées au prix de base.",
        ],
        "Évitez « formulaire standard » ou « moteur rapide ». Décrivez les champs, le destinataire, la confirmation, les erreurs et la preuve attendue.",
    ),
    SectionSpec(
        8,
        "Outils existants, intégrations et données",
        "OPTIONNEL · INTÉGRATIONS",
        "Une intégration découverte après le devis est une cause classique d'incertitude, de délai et de coût supplémentaire.",
        [],
        [
            "Confirmer la faisabilité seulement après accès à une documentation et, si possible, à un compte de test.",
            "Décrire les données échangées, la fréquence, la gestion des erreurs et la solution de secours.",
            "Isoler une phase d'étude si les inconnues empêchent un engagement ferme.",
        ],
        [
            "Chaque outil a un propriétaire interne et une finalité.",
            "Les inconnues sont visibles avant le chiffrage.",
            "Aucun secret ni identifiant réel ne figure dans le document.",
        ],
    ),
    SectionSpec(
        9,
        "Design, expérience et accessibilité",
        "ESSENTIEL · À ADAPTER AU PÉRIMÈTRE LÉGAL",
        "Des adjectifs comme « moderne » ou « intuitif » ne permettent pas d'évaluer une livraison ; des tâches et états observables, oui.",
        [
            Field("Identité disponible", "charte, logo, fichiers sources", "Charte légère et logo disponibles ; refonte complète de l'identité hors périmètre."),
            Field("À conserver", "repères de marque utiles", "Palette principale et ton sobre ; vérifier les contrastes avant réemploi."),
            Field("Perception recherchée", "effet et justification", "Crédible et concret : le lecteur doit comprendre les offres et les preuves sans superlatifs."),
            Field("Références appréciées", "URL et raison précise", "Deux références seront jointes pour leur hiérarchie et leur clarté, pas pour copier leur apparence."),
            Field("Appareils prioritaires", "contextes et tâches", "Mobile pour découverte/contact ; ordinateur pour comparaison et lecture de ressources."),
            Field("États à concevoir", "erreur, attente, absence", "Chargement, validation, erreur de champ, indisponibilité CRM, confirmation, absence de résultat."),
            Field("Périmètre légal", "analyse à faire", "Régime applicable à confirmer ; le service n'est pas un e-commerce dans cette V1."),
            Field("Cible d'accessibilité", "référentiel/tests/preuve", "Concevoir selon WCAG 2.2 niveau AA comme cible de qualité ; demander tests clavier, zoom, contraste, lecteurs d'écran sur parcours clés et rapport des écarts. Toute revendication de conformité exige un audit adapté."),
        ],
        [
            "Expliquer la méthode de conception, les états couverts et les tests prévus.",
            "Préciser ce qui est vérifié automatiquement et manuellement.",
            "Ne revendiquer une conformité RGAA ou WCAG qu'après l'audit approprié.",
        ],
        [
            "Les références sont commentées par raison, pas seulement listées.",
            "Les parcours au clavier, le zoom, les contrastes, les labels et messages d'erreur sont testables.",
            "Le champ légal est analysé pour le service concerné au lieu d'être présumé.",
        ],
    ),
    SectionSpec(
        10,
        "Référencement, mesure et migration",
        "REFONTE UNIQUEMENT · SEO",
        "Une refonte change des URL, des contenus et des signaux techniques ; le plan doit préserver l'information utile sans promettre un classement.",
        [],
        [
            "Fournir l'inventaire source, le plan ancienne URL vers nouvelle URL et le protocole de validation.",
            "Utiliser des redirections permanentes serveur 301 ou 308 vers la destination finale, sans chaînes évitables.",
            "Décrire les contrôles avant bascule et la surveillance des erreurs, du trafic et de l'indexation après lancement.",
        ],
        [
            "Chaque ancienne URL utile a une destination pertinente ou une décision justifiée.",
            "Les canonical, noindex, robots.txt, sitemap et propriétés Search Console sont contrôlés.",
            "Aucune garantie de position ni de délai d'indexation n'est formulée.",
        ],
        "Une migration peut provoquer des fluctuations temporaires. Les redirections utiles doivent être conservées au moins un an et plus longtemps lorsqu'elles continuent de servir les utilisateurs.",
    ),
    SectionSpec(
        11,
        "Données personnelles, traceurs et obligations",
        "SELON LE PROJET · À FAIRE VALIDER",
        "« Conforme RGPD » ne décrit ni les données, ni les rôles, ni les preuves nécessaires.",
        [],
        [
            "Lister les sous-traitants, lieux de traitement, transferts éventuels et garanties disponibles.",
            "Configurer les traceurs selon leurs finalités et recueillir le consentement avant ceux qui ne bénéficient pas d'une exemption.",
            "Séparer les responsabilités techniques de la validation des mentions, bases légales, durées et obligations sectorielles.",
        ],
        [
            "Les finalités, données, destinataires, durées et droits sont documentés ou marqués à valider.",
            "La présence ou l'absence de bandeau n'est pas décidée par une formule universelle.",
            "Les services sensibles, mineurs, paiement, profilage ou traitements massifs déclenchent une revue spécialisée.",
        ],
        "Un hébergement dans l'EEE ne suffit pas, à lui seul, à démontrer la conformité. La qualification des rôles, bases légales et obligations doit être validée pour le projet concerné.",
    ),
    SectionSpec(
        12,
        "Sécurité proportionnée aux risques",
        "ESSENTIEL · SÉCURITÉ",
        "La sécurité n'est pas un adjectif ni un état définitif ; elle repose sur des mesures, des responsables et un entretien.",
        [
            Field("Accès administratifs", "rôles, moindre privilège, MFA", "Comptes nominatifs, rôles distincts et MFA activée sur chaque compte privilégié lorsque le service le permet ; sinon écart justifié, mesure compensatoire et choix d'outil réexaminé. Aucun compte partagé."),
            Field("Secrets", "stockage et transmission", "Aucun secret dans le cahier des charges, les URL ou les données de test ; transmission séparée via un canal sécurisé."),
            Field("Environnements", "production, préproduction, test", "Préproduction protégée et non indexable ; données fictives pour les tests."),
            Field("Mises à jour", "responsable et délai", "Prestataire à préciser : périmètre maintenu, suivi des dépendances et délais selon gravité."),
            Field("Sauvegardes", "fréquence, rétention, restauration", "Politique à proposer ; test de restauration documenté avant mise en ligne puis périodiquement."),
            Field("Journalisation", "événements utiles et accès", "Échecs administratifs et incidents utiles, sans collecte disproportionnée ; durée à définir."),
            Field("Incident", "contact, notification, preuves", "Canal d'alerte, informations minimales, conservation des preuves et responsabilités à préciser au contrat."),
            Field("Spam et abus", "mesures proportionnées", "Protection des formulaires sans bloquer inutilement les utilisateurs ; solution accessible demandée."),
        ],
        [
            "Présenter les mesures incluses, les limites, la maintenance requise et les responsabilités.",
            "Décrire le processus de correction et d'alerte, sans promettre « 100 % sécurisé ».",
            "Identifier les éléments exigeant un audit ou un spécialiste externe.",
        ],
        [
            "Les rôles, secrets, sauvegardes, mises à jour et incidents ont un responsable.",
            "Un test de restauration et des données fictives de recette sont prévus.",
            "Les mesures sont adaptées au risque et non copiées mécaniquement.",
        ],
    ),
    SectionSpec(
        13,
        "Réponse technique attendue",
        "PRESTATAIRE RÉPOND",
        "Le client exprime ses contraintes ; le prestataire explique une solution, ses alternatives, ses limites et son coût d'exploitation.",
        [
            Field("Contrainte technique imposée", "uniquement si réellement nécessaire", "Aucune technologie imposée. Compatibilité avec le domaine, les comptes et le processus éditorial existants."),
            Field("Autonomie attendue", "tâches que l'équipe fera seule", "Modifier textes/images, créer une page depuis un modèle, publier une ressource et corriger les métadonnées après formation."),
            Field("Hébergement", "contraintes de localisation, compte, facturation", "Compte et facturation au nom de l'entreprise ; proposition et garanties à détailler."),
            Field("Compatibilité", "navigateurs/appareils représentatifs", "Deux dernières versions des principaux navigateurs et parcours clés sur mobile/ordinateur ; échantillon final à confirmer."),
            Field("Performance avant lancement", "pages, outil, réseau, répétitions", "Protocole laboratoire à proposer sur pages représentatives, appareil/réseau documentés, cache froid/chaud et plusieurs passes."),
            Field("Performance après lancement", "mesure réelle si données suffisantes", "Suivre LCP, INP et CLS au p75 lorsque le volume terrain devient suffisant ; repères actuels 2,5 s, 200 ms et 0,1."),
        ],
        [
            "Solution recommandée, justification et au moins une alternative examinée.",
            "Limites, dépendances, licences, services tiers et coûts récurrents.",
            "Architecture d'hébergement, réversibilité, équipe, sous-traitance et hypothèses.",
            "Protocole de performance avant lancement et dispositif de suivi réel après lancement.",
        ],
        [
            "La réponse permet de comprendre pourquoi la solution convient au besoin.",
            "Les coûts et dépendances à trois ans sont visibles.",
            "Les données terrain ne sont pas présentées comme disponibles dès la livraison.",
        ],
    ),
    SectionSpec(
        14,
        "Livrables, responsabilités et calendrier",
        "ESSENTIEL · GOUVERNANCE",
        "Un planning n'est crédible que si chaque livrable, validation, dépendance et retard possible a un responsable.",
        [],
        [
            "Présenter un calendrier par jalons, avec dépendances et délais de validation client.",
            "Distinguer les livrables inclus, les fichiers sources, la formation et la documentation.",
            "Décrire le processus de demande de changement et son impact avant exécution.",
        ],
        [
            "Chaque livrable a un responsable, une date et un décideur.",
            "Le nombre de cycles de retours et les délais de validation sont explicites.",
            "Les conséquences d'un retard ou d'une demande nouvelle sont traitées sans ambiguïté.",
        ],
    ),
    SectionSpec(
        15,
        "Recette, anomalies, garantie et acceptation",
        "ESSENTIEL · PREUVE",
        "Préparer la recette avant le développement transforme les attentes vagues en scénarios observables.",
        [
            Field("Environnement de recette", "URL, accès, données", "Préproduction protégée, données fictives et comptes de test remis séparément."),
            Field("Durée", "fenêtre et disponibilité", "Dix jours ouvrés prévus, à confirmer au contrat."),
            Field("Bloquant", "définition projet", "Empêche un parcours critique ou la mise en ligne sans solution de contournement acceptable."),
            Field("Majeur", "définition projet", "Dégrade fortement une fonction importante, avec ou sans contournement limité."),
            Field("Mineur", "définition projet", "Défaut non bloquant n'empêchant pas l'usage principal."),
            Field("Acceptation", "procédure et réserves", "Procédure à définir au contrat : procès-verbal, réserves, corrections, nouveau test et décision explicite."),
            Field("Garantie corrective", "durée, point de départ, exclusions", "Le prestataire doit proposer sa durée, son point de départ et distinguer défaut de livraison, maintenance et évolution."),
        ],
        [
            "Fournir le rapport de tests et les preuves convenues.",
            "Corriger puis soumettre au nouveau test selon le niveau d'anomalie et le contrat.",
            "Ne promettre ni « zéro bug » ni conformité globale sans protocole défini.",
        ],
        [
            "Chaque critère a une précondition, une action, un résultat attendu, un environnement et une preuve.",
            "Les niveaux d'anomalie et la procédure de réserve sont compris par les deux parties.",
            "Garantie corrective, maintenance et évolution sont séparées.",
        ],
        "Utilisez le fichier « grille-de-recette-site-internet.xlsx » du kit pour exécuter et tracer les tests.",
    ),
    SectionSpec(
        16,
        "Exploitation, comptes, droits et réversibilité",
        "ESSENTIEL · APRÈS LA MISE EN LIGNE",
        "Un site reste exploitable si l'entreprise contrôle ses comptes, connaît ses licences et peut récupérer ses contenus, données et documentation.",
        [],
        [
            "Distinguer code spécifique, composants open source, thèmes, extensions, polices, photos, contenus et services tiers.",
            "Proposer une cession ou licence adaptée pour les créations concernées, à formaliser au contrat.",
            "Décrire export, documentation, remise des accès, coût et délai de sortie.",
        ],
        [
            "Le titulaire du domaine, le compte d'hébergement, le dépôt et les administrateurs sont identifiés.",
            "Les droits transmis, leur périmètre et les licences tierces sont distingués.",
            "La procédure de sortie est réalisable sans dépendre d'un mot de passe partagé.",
        ],
        "Les droits ne se résument pas à « tout vous appartient ». En droit français, une cession doit notamment délimiter les droits transmis et leur domaine d'exploitation ; faites adapter le contrat.",
    ),
    SectionSpec(
        17,
        "Budget et format de l'offre",
        "ESSENTIEL · OFFRES COMPARABLES",
        "La comparaison devient utile lorsque les candidats répondent avec le même découpage, les mêmes hypothèses et un coût d'exploitation visible.",
        [
            Field("Enveloppe de construction", "HT/TTC et périmètre", "12 000 à 16 000 € HT — valeur entièrement fictive."),
            Field("Réserve interne", "séparée du budget annoncé", "10 % conservés en interne pour décisions nouvelles ; non inclus dans le périmètre de base."),
            Field("Coûts récurrents", "plafond ou attente", "Tous les services et licences annuels doivent être détaillés ; aucun plafond arbitraire avant proposition."),
            Field("Calendrier de paiement", "souhait ou contrainte", "Proposition attendue par jalons ; conditions à valider au contrat."),
            Field("Validité de l'offre", "durée attendue", "Au moins 30 jours souhaités, à confirmer par le candidat."),
            Field("Coût total", "horizon utile", "Présenter séparément investissement initial et coût estimatif sur trois ans, selon hypothèses explicites."),
        ],
        [
            "Ventiler cadrage, design, développement, contenus, migration, intégrations, recette, formation et mise en ligne.",
            "Détailler hébergement, licences, maintenance, options, coûts de sortie, hypothèses et exclusions.",
            "Marquer chaque exigence : incluse, partielle, optionnelle, alternative ou exclue.",
        ],
        [
            "Deux offres peuvent être rapprochées ligne par ligne.",
            "Le coût initial ne masque pas les dépenses récurrentes ou de sortie.",
            "Une option n'est pas présentée comme incluse dans le prix de base.",
        ],
    ),
    SectionSpec(
        18,
        "Questions, décisions et annexes",
        "ESSENTIEL · TRAÇABILITÉ",
        "Un registre commun réduit les réponses contradictoires et préserve l'équité entre candidats.",
        [
            Field("Canal de questions", "adresse ou outil dédié", "Canal fictif géré par la responsable communication ; aucune donnée sensible par email."),
            Field("Date limite", "dernière date de question", "20 août 2026 — date fictive."),
            Field("Mode de réponse", "à tous ou individuel", "Réponse consolidée et expurgée envoyée à tous les candidats sous forme d'addendum."),
            Field("Décisions ouvertes", "arbitrages restant à prendre", "Faisabilité CRM, conservation de trois contenus anciens et niveau de support après lancement."),
            Field("Annexes fournies", "liste sans secrets", "Arborescence, inventaire de 42 URL, statistiques agrégées, charte, exemples de contenus et grille de recette."),
        ],
        [
            "Lister les questions bloquantes et les hypothèses prises faute de réponse.",
            "Accuser réception de chaque addendum et l'intégrer à l'offre.",
            "Ne pas demander de données ou accès sensibles dans le dossier de consultation.",
        ],
        [
            "Toutes les décisions ouvertes ont un propriétaire et une date cible.",
            "Chaque candidat dispose des mêmes informations de référence.",
            "Les annexes sont nommées, datées et expurgées.",
        ],
    ),
]


def add_section_specific(doc: Document, spec: SectionSpec, filled: bool) -> None:
    if spec.fields:
        add_field_table(doc, spec.fields, filled)

    if spec.number == 4:
        rows = [
            ("OBJ-01", "Demandes qualifiées / mois", "7 — CRM, moyenne fictive 6 mois", "14 à M+9", "Direction"),
            ("OBJ-02", "URL utiles préservées", "Inventaire à consolider", "100 % décidées avant bascule", "Communication"),
            ("REC-01", "Publication autonome", "Dépendance au prestataire", "Modifier une page en < 15 min après formation", "Communication"),
        ]
        add_matrix(
            doc,
            ["ID", "OBJECTIF / CRITÈRE", "DÉPART & SOURCE", "CIBLE / HORIZON", "SUIVI"],
            rows if filled else [],
            [850, 2300, 2350, 2460, 1400],
            empty_rows=4 if not filled else 0,
        )
        add_callout(doc, "Lecture", "OBJ-01 est un résultat métier à suivre. REC-01 est un critère de livraison qui peut être testé pendant la recette.", "emerald")
    elif spec.number == 6:
        rows = [
            ("Accueil", "Orienter vers l'offre pertinente", "Réécriture", "Communication", "V1"),
            ("6 pages services", "Expliquer problèmes, méthode, preuve", "Reprise + réécriture", "Experts", "V1"),
            ("3 pages secteurs", "Répondre aux besoins par contexte", "Création", "Experts", "V1"),
            ("Références", "Rassurer avec cas autorisés", "Migration sélective", "Direction", "V1"),
            ("Ressources", "Permettre lecture et téléchargement", "Migration", "Communication", "V1"),
            ("Espace client", "Dépôt et suivi", "Non applicable", "—", "Hors V1"),
        ]
        add_matrix(doc, ["PAGE / LOT", "OBJECTIF", "TRAITEMENT", "RESPONSABLE", "STATUT"], rows if filled else [], [1900, 2600, 1900, 1560, 1400], empty_rows=6 if not filled else 0)
    elif spec.number == 7:
        rows = [
            ("F-01", "Formulaire qualifiant", "Indispensable", "Validation, erreur, spam, CRM indisponible", "Message reçu + email + trace"),
            ("F-02", "Gestion des contenus", "Indispensable", "Brouillon, publication, média trop lourd", "Modification autonome testée"),
            ("F-03", "Transmission CRM", "Indispensable si faisable", "API inconnue, doublon, indisponibilité", "Test de bout en bout + secours"),
            ("F-04", "Bibliothèque avancée", "Version ultérieure", "Volumes à confirmer", "Option séparée"),
        ]
        add_matrix(doc, ["ID", "BESOIN", "PRIORITÉ", "CAS LIMITES", "PREUVE ATTENDUE"], rows if filled else [], [900, 2050, 1700, 2200, 2510], empty_rows=5 if not filled else 0)
    elif spec.number == 8:
        rows = [
            ("CRM interne", "Demandes qualifiées", "Formulaire → CRM", "API à confirmer", "Email structuré + saisie manuelle"),
            ("Agenda", "Prise de rendez-vous", "Aucun échange en V1", "Non applicable", "Lien de contact"),
            ("Mesure d'audience", "Pilotage agrégé", "Événements nécessaires", "Configuration juridique à valider", "Mesure minimale"),
        ]
        add_matrix(doc, ["OUTIL", "FINALITÉ", "DONNÉES / SENS", "INCONNUE", "SECOURS"], rows if filled else [], [1600, 1800, 1950, 2050, 1960], empty_rows=4 if not filled else 0)
    elif spec.number == 10:
        rows = [
            ("42 URL actuelles", "Sitemap + Search Console + mesure", "Conserver / fusionner / supprimer", "Avant maquettes"),
            ("Plan de redirections", "Ancienne → nouvelle destination finale", "301 ou 308 serveur", "Avant bascule"),
            ("Préproduction", "Noindex + accès restreint", "Contrôle robots/canonical", "Avant ouverture"),
            ("Après lancement", "404, 5xx, trafic, indexation", "J+1, J+7, J+30, puis suivi", "Responsables à nommer"),
        ]
        add_matrix(doc, ["LOT", "SOURCE / ACTION", "PREUVE", "MOMENT"], rows if filled else [], [1700, 3100, 2700, 1860], empty_rows=4 if not filled else 0)
    elif spec.number == 11:
        rows = [
            ("Formulaire contact", "Qualifier et répondre", "Coordonnées pro + besoin", "Équipe commerciale / CRM", "À valider"),
            ("Mesure d'audience", "Comprendre les parcours", "Données techniques agrégées", "Outil à choisir", "Exemption ou consentement à valider"),
            ("Newsletter", "Non incluse en V1", "Aucune", "—", "Hors périmètre"),
        ]
        add_matrix(doc, ["TRAITEMENT", "FINALITÉ", "DONNÉES", "DESTINATAIRE / OUTIL", "VALIDATION"], rows if filled else [], [1500, 1900, 1900, 2100, 1960], empty_rows=4 if not filled else 0)
    elif spec.number == 14:
        rows = [
            ("Cadrage", "Périmètre + planning", "Prestataire", "Direction", "5 jours ouvrés"),
            ("Conception", "Arborescence + maquettes + états", "Prestataire", "Communication", "5 jours / 2 cycles"),
            ("Contenus", "Textes et médias validés", "Client", "Direction", "Selon rétroplanning"),
            ("Développement", "Site + configuration", "Prestataire", "Communication", "Démonstrations jalonnées"),
            ("Recette", "Rapport + corrections", "Partagé", "Direction", "10 jours ouvrés"),
            ("Mise en ligne", "Accès, documentation, formation", "Prestataire", "Direction", "PV selon contrat"),
        ]
        add_matrix(doc, ["PHASE", "LIVRABLE", "RESPONSABLE", "VALIDATEUR", "DÉLAI / CYCLES"], rows if filled else [], [1300, 2500, 1700, 1600, 2260], empty_rows=6 if not filled else 0)
    elif spec.number == 16:
        rows = [
            ("Nom de domaine", "Entreprise", "Compte au nom de l'entreprise", "Transfert documenté"),
            ("Hébergement", "Entreprise / à confirmer", "Facturation directe souhaitée", "Export + documentation"),
            ("Dépôt de code", "À définir au contrat", "Accès administrateur client", "Historique et guide de reprise"),
            ("Contenus/données", "Entreprise", "Exports standards", "Export testé avant sortie"),
            ("Composants tiers", "Selon licences", "Inventaire obligatoire", "Conditions de remplacement"),
        ]
        add_matrix(doc, ["ACTIF", "TITULAIRE / DROITS", "ACCÈS ATTENDU", "SORTIE"], rows if filled else [], [1800, 2500, 2600, 2460], empty_rows=5 if not filled else 0)


def add_sources_and_license(doc: Document, filled: bool) -> None:
    # The editable model ends its final questionnaire close to the bottom of a
    # page, so its references need a deliberate clean page. The filled example
    # has enough room after its final checklist: letting the references flow
    # there avoids creating an almost-empty penultimate page.
    if not filled:
        doc.add_page_break()
    doc.add_heading("Sources, limites et licence d'utilisation", level=1)
    add_callout(
        doc,
        "Important",
        "Ce kit aide à structurer une consultation. Il ne remplace pas le devis, le contrat, un audit de conformité, une analyse de sécurité ni un conseil juridique personnalisé.",
        "amber",
    )
    doc.add_heading(
        f"Références officielles consultées le {PUBLICATION_DATE_LABEL}", level=2
    )
    sources = [
        ("France Num — bâtir le cahier des charges d'un site", "https://www.francenum.gouv.fr/guides-et-conseils/developpement-commercial/site-web/batir-le-cahier-des-charges-du-site-internet"),
        ("CNIL — cookies et autres traceurs", "https://www.cnil.fr/fr/cookies-et-autres-traceurs/regles/cookies/comment-mettre-mon-site-web-en-conformite"),
        ("CNIL — qualification juridique de la sous-traitance", "https://www.cnil.fr/fr/qualification-juridique-sous-traitance"),
        ("RGAA — champ d'application", "https://accessibilite.numerique.gouv.fr/obligations/champ-application/"),
        ("DGCCRF — directive européenne sur l'accessibilité", "https://www.economie.gouv.fr/dgccrf/les-fiches-pratiques/la-nouvelle-directive-europeenne-accessibilite-pour-des-produits-et-des-services-accessibles-aux-personnes-en-situation"),
        ("Google Search Central — Core Web Vitals", "https://developers.google.com/search/docs/appearance/core-web-vitals"),
        ("web.dev — évaluation au 75e percentile", "https://web.dev/articles/vitals?hl=fr"),
        ("Google Search Central — redirections 301/308", "https://developers.google.com/search/docs/crawling-indexing/301-redirects"),
        ("Google Search Central — déplacement de site et durée des redirections", "https://developers.google.com/search/docs/crawling-indexing/site-move-with-url-changes?hl=fr"),
        ("Légifrance — Code de la propriété intellectuelle, article L131-3", "https://www.legifrance.gouv.fr/codes/article_lc/LEGIARTI000006278958"),
        ("Afnic — guide du titulaire d'un nom de domaine", "https://www.afnic.fr/wp-media/uploads/2020/12/Guidepratique_Titulaire_VF.pdf"),
        ("MesServicesCyber — authentification multifacteur et mots de passe", "https://messervices.cyber.gouv.fr/guides/recommandations-relatives-lauthentification-multifacteur-et-aux-mots-de-passe"),
        ("Cybermalveillance.gouv.fr — bonnes pratiques de sauvegarde", "https://www.cybermalveillance.gouv.fr/tous-nos-contenus/bonnes-pratiques/sauvegardes"),
    ]
    add_link_bullets(doc, sources)
    doc.add_heading("Licence d'utilisation du kit", level=2)
    add_body(
        doc,
        "Vous pouvez utiliser et modifier ce document pour cadrer vos propres projets, le partager en interne et avec les prestataires consultés. La revente ou republication du kit, complet ou substantiellement identique, comme ressource autonome est interdite sans accord écrit de Hagnéré Code. Vos réponses, ajouts et données restent sous votre responsabilité.",
    )
    add_body(
        doc,
        f"Version {VERSION} — {PUBLICATION_DATE_LABEL}. Pour signaler une erreur : utilisez la page de contact de hagnere-code.ai sans joindre de cahier des charges contenant des données sensibles.",
    )
    # The fictitious nature of the filled example is already stated on its
    # cover and before the sources. Repeating it here would push a single
    # callout onto an otherwise empty final page after PDF conversion.


def build_docx(output: Path, filled: bool) -> None:
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    set_doc_language(doc)
    configure_properties(
        doc,
        "Exemple fictif rempli — cahier des charges de site internet"
        if filled
        else "Modèle de cahier des charges de site internet",
        "Kit opérationnel pour cadrer, consulter des prestataires et préparer la recette.",
    )
    add_header_footer(doc, filled)
    cover(doc, filled)
    add_usage(doc, filled)
    for spec in SECTIONS:
        add_section_intro(doc, spec)
        add_section_specific(doc, spec, filled)
        add_provider_and_done(doc, spec)
    add_sources_and_license(doc, filled)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_readme_docx(output: Path) -> None:
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    set_doc_language(doc)
    configure_properties(
        doc,
        "Mode d'emploi du kit cahier des charges de site internet",
        "Utiliser le modèle Word, l'exemple rempli et la grille de recette.",
    )
    add_header_footer(doc, filled=False, label_override="MODE D'EMPLOI · ")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    add_tag(p, "RESSOURCE OPÉRATIONNELLE", VIOLET)
    title = doc.add_paragraph(style="Title")
    title.add_run("Kit cahier des charges de site internet")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Mode d'emploi court · modèle Word, exemple rempli et grille de recette Excel")
    add_callout(
        doc,
        "Téléchargement",
        f"Le kit est proposé sans formulaire ni adresse email obligatoire. Version {VERSION} — {PUBLICATION_DATE_LABEL}.",
        "emerald",
    )
    doc.add_heading("Ce que contient le kit", level=1)
    add_matrix(
        doc,
        ["FICHIER", "UTILITÉ", "FORMAT"],
        [
            ("Modèle", "Cadrer le besoin, les exclusions, les responsabilités et la réponse attendue.", "DOCX éditable"),
            ("Exemple", "Voir un dossier fictif rempli de bout en bout sans faux résultat client.", "PDF"),
            ("Grille de recette", "56 tests préremplis, 12 lignes libres, preuves, anomalies et synthèse.", "XLSX"),
            ("Mode d'emploi", "Choisir le bon ordre et éviter les erreurs de partage.", "PDF"),
        ],
        [1900, 5660, 1800],
    )
    doc.add_heading("Résultat attendu", level=2)
    add_body(
        doc,
        "Deux prestataires doivent pouvoir comprendre le même périmètre, les mêmes exclusions et les mêmes preuves de livraison. Le kit peut aussi conduire à réduire le projet si un site plus simple suffit.",
    )
    doc.add_page_break()

    doc.add_heading("Le parcours le plus rapide", level=1)
    add_bullets(
        doc,
        [
            "10 min — Synthèse : problème, résultat principal, public, indispensable, hors périmètre, budget et date.",
            "20 min — Périmètre : pages, contenus, fonctionnalités et outils à connecter.",
            "15 min — Après lancement : comptes, maintenance, droits, coûts récurrents et sortie.",
            "15 min — Recette : choisir les parcours critiques et les preuves attendues.",
            "10 min — Nettoyage : supprimer les modules inutiles, vérifier les secrets, exporter une copie PDF et figer la version.",
        ],
        numbered=True,
    )
    doc.add_heading("Trois catégories de réponses", level=2)
    add_matrix(
        doc,
        ["CATÉGORIE", "CE QU'ELLE CONTIENT"],
        [
            ("Vous décidez", "Besoin, priorités, budget, exclusions et responsables."),
            ("Le prestataire propose", "Solution, alternatives, limites, planning détaillé et coûts."),
            ("À faire valider", "Contrat, droits, RGPD, accessibilité applicable, sécurité sensible ou obligations sectorielles."),
        ],
        [2500, 6860],
    )
    doc.add_heading("Avant d'envoyer", level=2)
    add_bullets(
        doc,
        [
            "Exporter une copie PDF identique pour tous les candidats et conserver le Word comme source interne.",
            "Retirer commentaires, historique de révision, métadonnées personnelles et données sensibles.",
            "Ne jamais insérer mot de passe, clé API, accès administrateur, fichier client brut ou donnée sensible.",
            "Envoyer les questions et réponses communes à tous sous forme d'addendum daté.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("Ce que le kit évite volontairement", level=1)
    add_body(
        doc,
        "Les expressions suivantes semblent rassurantes mais ne décrivent pas une livraison vérifiable. Le modèle les transforme en décisions, scénarios et preuves.",
    )
    add_matrix(
        doc,
        ["À ÉVITER", "À PRÉCISER"],
        [
            ("Site moderne et intuitif", "Tâches, hiérarchie, états, erreurs et tests utilisateurs."),
            ("Conforme RGPD", "Données, finalités, rôles, sous-traitants, durées, preuve et validation."),
            ("Accessible", "Périmètre légal, cible, référentiel, tests manuels et rapport."),
            ("Rapide", "Pages, outil, appareil, réseau, cache, répétitions et mesure terrain distincte."),
            ("Optimisé SEO", "Inventaire, contenus, métadonnées, redirections, contrôles et suivi."),
            ("Le code vous appartient", "Droits transmis, licence ou cession, composants tiers, comptes et réversibilité."),
            ("Zéro bug", "Scénarios, gravité, preuve, réserve, correction et nouveau test."),
        ],
        [2700, 6660],
    )
    add_callout(
        doc,
        "À propos de l'exemple",
        "L'entreprise, les personnes, les chiffres, les URL, les budgets, le CRM, le calendrier et les résultats sont inventés. L'exemple n'est ni une référence client ni une preuve de performance.",
        "gray",
    )
    doc.add_page_break()

    doc.add_heading("Limites, maintenance et licence", level=1)
    add_callout(
        doc,
        "Limite",
        "Ce kit aide à cadrer un projet. Il ne remplace pas le devis, le contrat, un audit de conformité, une analyse de sécurité ou un conseil juridique personnalisé. Une revendication de conformité nécessite les vérifications ou audits appropriés.",
        "amber",
    )
    doc.add_heading("Références officielles", level=2)
    add_link_bullets(
        doc,
        [
            ("France Num — bâtir un cahier des charges", "https://www.francenum.gouv.fr/guides-et-conseils/developpement-commercial/site-web/batir-le-cahier-des-charges-du-site-internet"),
            ("CNIL — cookies et traceurs", "https://www.cnil.fr/fr/cookies-et-autres-traceurs/regles/cookies/comment-mettre-mon-site-web-en-conformite"),
            ("RGAA — champ d'application", "https://accessibilite.numerique.gouv.fr/obligations/champ-application/"),
            ("Google Search Central — Core Web Vitals", "https://developers.google.com/search/docs/appearance/core-web-vitals"),
            ("Google Search Central — déplacement de site", "https://developers.google.com/search/docs/crawling-indexing/site-move-with-url-changes?hl=fr"),
            ("Légifrance — article L131-3", "https://www.legifrance.gouv.fr/codes/article_lc/LEGIARTI000006278958"),
            ("MesServicesCyber — authentification", "https://messervices.cyber.gouv.fr/guides/recommandations-relatives-lauthentification-multifacteur-et-aux-mots-de-passe"),
        ],
    )
    doc.add_heading("Licence d'utilisation", level=2)
    add_body(
        doc,
        "Vous pouvez utiliser et modifier le kit pour vos propres projets, le partager en interne et avec les prestataires consultés. Sa revente ou republication, complète ou substantiellement identique, comme ressource autonome est interdite sans accord écrit de Hagnéré Code.",
    )
    add_body(
        doc,
        "Pour signaler une erreur, utilisez la page de contact de hagnere-code.ai sans joindre de document contenant des données sensibles.",
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    out = args.output_dir
    build_docx(out / "modele-cahier-des-charges-site-internet.docx", filled=False)
    build_docx(out / "exemple-rempli-cahier-des-charges-site-internet-source.docx", filled=True)
    build_readme_docx(out / "lisez-moi-kit-cahier-des-charges-site-internet-source.docx")
    print(f"Built document assets in {out}")


if __name__ == "__main__":
    main()
